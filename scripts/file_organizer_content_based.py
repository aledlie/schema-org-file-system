#!/usr/bin/env python3
"""
Content-Based Intelligent File Organizer using Schema.org metadata and OCR.

Organizes files based on their actual content rather than just file type.
Uses OCR to extract text from images and PDFs, then classifies by content.
"""

import json
import os
import re
import shutil
import sys
from collections import defaultdict
from contextlib import nullcontext
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import quote

# Shared filename-pattern classifier (single source of truth, also used by
# src/organizers/content_organizer.py). Re-export research helpers that the rest
# of this module and tests still reference by their original names.
from shared.filename_classifier import (  # noqa: E402,F401  (re-exported for tests)
    RESEARCH_CATEGORY,
    SCHOLARLY_ARTICLE_SCHEMA_TYPE,
    _detect_research_publisher,
)
from shared.filename_classifier import (  # noqa: E402
    classify_by_filename_patterns as _classify_by_filename_patterns,
)

# OCR (docTR via shared.ocr_classifier) and PDF imports
try:
    import pypdf
    from PIL import Image
    from shared.file_ops import resolve_collision
    from shared.filename_utils import is_generic_filename
    from shared.ocr_classifier import SCREENSHOT_KEYWORDS
    from shared.ocr_classifier import classify_by_ocr as _shared_classify_by_ocr
    from shared.ocr_classifier import (
        extract_ocr_text,
        extract_ocr_text_pdf,
        extract_ocr_with_confidence,
        is_ocr_available,
    )
    from shared.status import ProcessingStatus

    OCR_AVAILABLE = is_ocr_available()

    # HEIC support
    try:
        from pillow_heif import register_heif_opener

        register_heif_opener()
    except ImportError:
        pass
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not available. Install python-doctr[torch], Pillow, pypdf")

# KIE (Key Information Extraction) imports
try:
    from shared.kie_schema_mapping import kie_result_to_schema_org
    from shared.kie_utils import extract_kie_fields, is_kie_available

    KIE_AVAILABLE = is_kie_available()
except ImportError:
    KIE_AVAILABLE = False

# Word document imports
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install python-docx")

# Excel imports
try:
    from openpyxl import load_workbook

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Warning: openpyxl not available. Install openpyxl")

# Add src directory to path (portable)
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from rename_images import IMAGE_EXTENSIONS_WIDE, PHOTO_PROFILE, ImageAnalyzer  # noqa: E402

from analyzers.image_analyzer import ImageContentAnalyzer  # noqa: E402
from base import PropertyType  # noqa: E402
from enrichment import MetadataEnricher, cached_stat  # noqa: E402
from generators import DocumentGenerator, ImageGenerator  # noqa: E402
from integration import SchemaRegistry  # noqa: E402
from validator import SchemaValidator  # noqa: E402

# Graph storage imports
try:
    from storage.graph_store import GraphStore
    from storage.models import FileStatus

    GRAPH_STORE_AVAILABLE = True
except ImportError:
    GRAPH_STORE_AVAILABLE = False
    print("Warning: GraphStore not available. Database persistence disabled.")

# CLIP classifier — used by the weak-image enhancement signal (_run_clip_signal).
# Image composition/face detection lives in analyzers.ImageContentAnalyzer.
try:
    from shared.clip_utils import get_clip_classifier
except ImportError:
    get_clip_classifier = None
    print("Warning: Vision libraries not available. Install open-clip-torch, torch, opencv-python")

# Image metadata imports
try:
    from geopy.exc import GeocoderServiceError, GeocoderTimedOut
    from geopy.geocoders import Nominatim
    from PIL import Image  # noqa: F811 - re-imported for the metadata path's own availability guard
    from PIL.ExifTags import GPSTAGS, TAGS

    METADATA_AVAILABLE = True

    # HEIC support
    try:
        from pillow_heif import register_heif_opener

        register_heif_opener()
        HEIC_AVAILABLE = True
    except ImportError:
        HEIC_AVAILABLE = False
except ImportError:
    METADATA_AVAILABLE = False
    print("Warning: Metadata libraries not available. Install piexif, geopy")

# Cost tracking imports (optional - gracefully degrade if not available)
try:
    from cost_roi_calculator import CostROICalculator, CostTracker

    COST_TRACKING_AVAILABLE = True
except ImportError:
    COST_TRACKING_AVAILABLE = False

    # Provide stub implementations for graceful degradation
    class CostTracker:
        """Stub CostTracker when cost tracking is not available."""

        def __init__(self, *args, **kwargs):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False


# Error tracking imports (optional - gracefully degrade if not available)
try:
    from error_tracking import (
        ErrorLevel,
        FileProcessingErrorTracker,
        capture_error,
        init_sentry,
        track_error,
        track_operation,
    )

    ERROR_TRACKING_AVAILABLE = True
except ImportError:
    ERROR_TRACKING_AVAILABLE = False

    # Stub implementations
    def init_sentry(*args, **kwargs):
        return False

    def capture_error(*args, **kwargs):
        pass

    def track_operation(*args, **kwargs):
        from contextlib import nullcontext

        return nullcontext()

    def track_error(*args, **kwargs):
        def decorator(func):
            return func

        return decorator

    class FileProcessingErrorTracker:
        def __init__(self):
            pass

        def track_file(self, *args, **kwargs):
            from contextlib import nullcontext

            return nullcontext()

        def print_summary(self):
            pass

        def get_stats(self):
            return {}

    class ErrorLevel:
        FATAL = "fatal"
        ERROR = "error"
        WARNING = "warning"
        INFO = "info"
        DEBUG = "debug"


# CLIP enhancement constants for weak image classification
try:
    from shared.constants import (
        CLIP_CATEGORY_PROMPTS,
        CLIP_CONTENT_LABELS,
        CLIP_ENHANCE_THRESHOLD,
        CLIP_LABEL_TO_ORGANIZER,
    )

    ENHANCED_CLIP_AVAILABLE = True
except ImportError:
    ENHANCED_CLIP_AVAILABLE = False

# CLIP cache support
try:
    from shared.clip_cache import CLIP_CACHE_AVAILABLE, get_cached_embedding
except ImportError:
    CLIP_CACHE_AVAILABLE = False

# Reverse map: full CLIP prompt → short CLIP_CONTENT_LABELS key.
# Built once at import time; used by enhance_weak_image_classification.
_CLIP_PROMPT_TO_LABEL: dict = (
    dict(zip(CLIP_CATEGORY_PROMPTS, CLIP_CONTENT_LABELS)) if ENHANCED_CLIP_AVAILABLE else {}
)

# Minimum OCR confidence required to treat text as reliable for ID doc detection.
# Low-confidence OCR on a photo (e.g. a blurry selfie) must not trigger passport detection.
# Uses docTR word-level confidence (true 0–1 scale).
_OCR_CONFIDENCE_THRESHOLD = 0.3

# Minimum keyword-hit ratio to accept a screenshot OCR sub-classification.
# classify_by_ocr() scores as hits/len(keywords), so this is calibrated to that
# scale — NOT to CLIP probability space.  With SCREENSHOT_MIN_HITS=2 and the
# largest category having 14 keywords, 2 hits = 14.3%.  A threshold of 0.10
# accepts any category that cleared SCREENSHOT_MIN_HITS (the real quality bar);
# the 0.30 gate previously used was copied from _OCR_CONFIDENCE_THRESHOLD and
# silently rejected valid 3–18% scores as "low confidence".
_SCREENSHOT_OCR_KEYWORD_THRESHOLD = 0.10

# Unified-scoring weights for the image classification path.
# CLIP score (already 0-1) is used as-is; OCR-text contributes a prior scaled by
# extraction length, plus an agreement boost when both signals point at the same
# (category, subcategory). Values are tuned so that text — which is typically
# more semantically specific than CLIP — wins ties on longer OCR extractions,
# while CLIP retains primacy when text extraction is sparse or unparseable.
_TEXT_SIGNAL_PRIOR = 0.80
_TEXT_LENGTH_FULL_CHARS = 200
_TEXT_MIN_CHARS = 30
_SIGNAL_AGREEMENT_BOOST = 0.15

# Personal titles that strongly indicate a human (vs. an org/brand name).
_HUMAN_TITLE_RE = re.compile(
    r"\b(?:Mr|Mrs|Ms|Miss|Dr|Prof|Sir|Madam|Hon|Rev|Esq|Atty)\.?\s+[A-Z]",
)
# First-person / signatory phrases that only humans write about themselves.
_HUMAN_CONTACT_PHRASES = (
    "date of birth",
    "d.o.b",
    "dob:",
    "signed by",
    "signature of",
    "undersigned",
    "to whom it may concern",
    "i hereby",
    "i am pleased to",
    "social security",
    "ssn:",
    "driver license",
    "driver's license",
    "maiden name",
    "next of kin",
    "emergency contact",
)


def _has_human_name_signal(text: str) -> bool:
    """
    Require evidence that a document is about a human, not an org/brand.

    Org-precedence rule: when none of these signals appear, defer person
    classification so org/document-type classifiers can win on names like
    "Morning Train" that look human but aren't.
    """
    if _HUMAN_TITLE_RE.search(text):
        return True
    text_lower = text.lower()
    return any(phrase in text_lower for phrase in _HUMAN_CONTACT_PHRASES)


class ContentClassifier:
    """Classifies document content into categories."""

    def __init__(self):
        """Initialize classifier with keyword patterns."""
        # Company name patterns
        self.company_patterns = [
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+LLC\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+L\.L\.C\.\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Inc\.?\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Incorporated\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Corp\.?\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Corporation\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Company\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Co\.\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Ltd\.?\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Limited\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+LLP\b",
            r"\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+L\.L\.P\.\b",
        ]

        # People name patterns - look for common name patterns
        self.people_patterns = [
            # ALL-CAPS names at start of resume (common in templates)
            # Matches: "ISABEL BUDENZ\nLLM" or "JOHN DOE\nSoftware Engineer"
            r"^([A-Z]{2,})\s+([A-Z]{2,})\s*\n",
            # ALL-CAPS name followed by title/degree
            r"\b([A-Z]{2,})\s+([A-Z]{2,})\s*\n\s*(?:LLM|MBA|PhD|MD|JD|CPA|Software|Engineer|Manager|Director|Analyst)",  # noqa: E501
            # Name with document type indicators
            r"\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\s+(?:Resume|CV|Cover Letter)\b",
            r"\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\s+(?:Portfolio|Biography|Bio)\b",
            # Field labels followed by names
            r"\b(?:Name|Contact|From|To|Attn|Author|Client|Patient|Student):\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b",  # noqa: E501
            # Email signatures (name before email)
            r"\b([A-Z][a-z]+\s+[A-Z][a-z]+)\s+<[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}>",
            # Name in "Prepared by/for" statements
            r"\b(?:Prepared|Written|Submitted|Signed)\s+(?:by|for):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b",  # noqa: E501
            # Name followed by credentials (MD, PhD, Esq, etc.)
            r"\b([A-Z][a-z]+\s+[A-Z][a-z]+),?\s+(?:MD|PhD|Esq|DDS|CPA|MBA|JD|RN)\b",
            # Mr./Mrs./Ms./Dr. followed by name
            r"\b(?:Mr|Mrs|Ms|Dr|Prof)\.?\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b",
            # Name in meeting notes format
            r"\b(?:Attendee|Participant|Speaker|Presenter):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b",
        ]

        self.patterns = {
            "legal": {
                "keywords": [
                    "contract",
                    "agreement",
                    "terms",
                    "conditions",
                    "legal",
                    "attorney",
                    "law",
                    "litigation",
                    "plaintiff",
                    "defendant",
                    "court",
                    "settlement",
                    "lease",
                    "deed",
                    "will",
                    "testament",
                    "power of attorney",
                    "notary",
                    "amendment",
                    "exhibit",
                    "whereas",
                    "party",
                    "parties",
                    "executed",
                    "operating agreement",
                    "llc",
                    "corporation",
                    "bylaws",
                    "articles",
                ],
                "subcategories": {
                    "contracts": ["contract", "agreement", "terms", "subscription", "saas"],
                    "real_estate": ["lease", "deed", "property", "real estate", "mortgage"],
                    "corporate": [
                        "llc",
                        "corporation",
                        "operating agreement",
                        "bylaws",
                        "articles",
                        "formation",
                    ],
                    "other": [],
                },
            },
            "financial": {
                "keywords": [
                    "invoice",
                    "receipt",
                    "tax",
                    "irs",
                    "payment",
                    "bill",
                    "billing",
                    "statement",
                    "account",
                    "balance",
                    "transaction",
                    "credit",
                    "debit",
                    "bank",
                    "finance",
                    "loan",
                    "interest",
                    "1098",
                    "1099",
                    "w-2",
                    "w2",
                    "federal",
                    "state return",
                    "refund",
                    "revenue",
                    "expense",
                    "budget",
                    "investment",
                    "portfolio",
                    "ein",
                    "employer identification",
                ],
                "subcategories": {
                    "tax": ["tax", "irs", "1098", "1099", "w-2", "w2", "federal", "state return"],
                    "invoices": ["invoice", "bill", "billing", "payment"],
                    "statements": ["statement", "account", "balance", "transaction"],
                    "other": [],
                },
            },
            "business": {
                "keywords": [
                    "proposal",
                    "pitch",
                    "business plan",
                    "strategy",
                    "marketing",
                    "presentation",
                    "deck",
                    "startup",
                    "company",
                    "venture",
                    "investor",
                    "growth",
                    "revenue model",
                    "unit economics",
                    "expansion",
                    "rfp",
                    "guidelines",
                    "program",
                    "service package",
                    "pricing",
                    "client",
                    "customer",
                    "vendor",
                    "supplier",
                    "partner",
                    "contacts",
                    "crm",
                    "hiring",
                    "job posting",
                    "meeting",
                    "standup",
                    "minutes",
                ],
                "subcategories": {
                    "planning": ["business plan", "strategy", "expansion", "growth", "project"],
                    "marketing": ["marketing", "pricing", "service package", "pitch", "deck"],
                    "proposals": ["proposal", "rfp", "guidelines"],
                    "crm": ["crm", "contacts", "microlender", "customer"],
                    "hr": ["hiring", "job posting", "team roster", "application", "linkedin"],
                    "meeting_notes": ["meeting", "standup", "minutes", "agenda", "retrospective"],
                    "clients": ["client", "llc", "inc", "corp", "company"],  # Legacy
                    "other": [],
                },
            },
            "personal": {
                "keywords": [
                    "resume",
                    "cv",
                    "cover letter",
                    "curriculum vitae",
                    "employment",
                    "personal",
                    "identification",
                    "passport",
                    "driver license",
                    "ssn",
                    "birth certificate",
                    "marriage",
                    "divorce",
                    "diploma",
                    "transcript",
                    "reference",
                    "recommendation",
                ],
                "subcategories": {
                    "employment": ["resume", "cv", "cover letter", "employment", "reference"],
                    "identification": ["passport", "driver license", "ssn", "id"],
                    "certificates": ["birth certificate", "marriage", "divorce", "diploma"],
                    "other": [],
                },
            },
            "medical": {
                "keywords": [
                    "medical",
                    "health",
                    "doctor",
                    "patient",
                    "prescription",
                    "diagnosis",
                    "treatment",
                    "hospital",
                    "clinic",
                    "insurance claim",
                    "hipaa",
                    "vaccination",
                    "immunization",
                    "lab results",
                    "pharmacy",
                ],
                "subcategories": {
                    "records": ["medical record", "patient", "diagnosis", "treatment"],
                    "insurance": ["insurance", "claim", "coverage"],
                    "prescriptions": ["prescription", "pharmacy", "medication"],
                    "other": [],
                },
            },
            "property": {
                "keywords": [
                    "property management",
                    "tenant",
                    "landlord",
                    "rent",
                    "rental",
                    "maintenance",
                    "repair",
                    "inspection",
                    "utilities",
                    "hoa",
                ],
                "subcategories": {
                    "leases": ["lease", "tenant", "landlord", "rent", "rental"],
                    "maintenance": ["maintenance", "repair", "inspection"],
                    "other": [],
                },
            },
            "education": {
                "keywords": [
                    "course",
                    "syllabus",
                    "lecture",
                    "assignment",
                    "homework",
                    "exam",
                    "grade",
                    "transcript",
                    "diploma",
                    "degree",
                    "certificate",
                    "university",
                    "college",
                    "school",
                    "research paper",
                    "thesis",
                    "dissertation",
                ],
                "subcategories": {
                    "coursework": ["course", "syllabus", "lecture", "assignment"],
                    "research": ["research", "paper", "thesis", "dissertation"],
                    "records": ["transcript", "diploma", "degree", "certificate"],
                    "other": [],
                },
            },
            "technical": {
                "keywords": [
                    "code",
                    "software",
                    "development",
                    "programming",
                    "api",
                    "database",
                    "documentation",
                    "technical",
                    "specification",
                    "architecture",
                    "design",
                    "system",
                    "infrastructure",
                    "deployment",
                    "configuration",
                ],
                "subcategories": {
                    "documentation": ["documentation", "spec", "specification", "readme"],
                    "architecture": ["architecture", "design", "system", "infrastructure"],
                    "other": [],
                },
            },
            "creative": {
                "keywords": [
                    "design",
                    "graphic",
                    "illustration",
                    "artwork",
                    "photo",
                    "image",
                    "screenshot",
                    "mockup",
                    "prototype",
                    "wireframe",
                    "brand",
                    "logo",
                ],
                "subcategories": {
                    "design": ["design", "mockup", "wireframe", "prototype"],
                    "branding": ["brand", "logo", "identity"],
                    "photos": ["photo", "photography", "image"],
                    "other": [],
                },
            },
        }

    def extract_company_names(self, text: str) -> List[str]:
        """
        Extract company names from text using regex patterns.

        Returns:
            List of detected company names
        """
        companies = []
        for pattern in self.company_patterns:
            matches = re.findall(pattern, text)
            companies.extend(matches)

        # Remove duplicates and clean up
        unique_companies = []
        seen = set()
        for company in companies:
            # Clean up whitespace
            clean = " ".join(company.split())
            # Skip if too short or already seen
            if len(clean) > 2 and clean.lower() not in seen:
                seen.add(clean.lower())
                unique_companies.append(clean)

        return unique_companies

    def _collapse_spaced_text(self, text: str) -> str:
        """
        Collapse spaced-out text like "I S A B E L  B U D E N Z" to "ISABEL BUDENZ".
        Common in stylized resume/CV templates.
        """

        # Pattern: single letters separated by spaces (at least 3 in a row)
        # Match sequences like "I S A B E L" (single chars with single spaces)
        def collapse_match(match):
            spaced = match.group(0)
            # Remove single spaces between single characters
            collapsed = re.sub(r"(?<=\b[A-Z]) (?=[A-Z]\b)", "", spaced)
            return collapsed

        # Find sequences of spaced single uppercase letters
        # Pattern matches: capital letter, space, capital letter (repeated)
        result = re.sub(r"\b([A-Z] ){2,}[A-Z]\b", collapse_match, text)
        return result

    def extract_people_names(self, text: str) -> List[str]:
        """
        Extract people names from text using regex patterns.

        Returns:
            List of detected people names
        """
        # Preprocess: collapse spaced-out text (common in stylized resumes)
        text = self._collapse_spaced_text(text)

        people = []
        for pattern in self.people_patterns:
            matches = re.findall(pattern, text)
            # Pattern can return tuples (first, last) or single strings
            for match in matches:
                if isinstance(match, tuple):
                    # Join tuple elements (e.g., first name + last name)
                    full_name = " ".join([m for m in match if m])
                else:
                    full_name = match
                people.append(full_name)

        # Remove duplicates and clean up
        unique_people = []
        seen = set()
        for person in people:
            # Clean up whitespace
            clean = " ".join(person.split())
            # Convert ALL-CAPS to Title Case (common in resume headers)
            if clean.isupper():
                clean = clean.title()
            # Skip if too short or already seen
            if len(clean) > 2 and clean.lower() not in seen:
                seen.add(clean.lower())
                unique_people.append(clean)

        return unique_people

    def extract_person_company_relationships(self, text: str) -> Dict[str, str]:
        """
        Extract relationships between people and companies from text.
        Uses Schema.org-style connections (Person worksFor/memberOf Organization).

        Returns:
            Dictionary mapping person names to company names
        """
        relationships = {}

        # Patterns for person-company relationships
        relationship_patterns = [
            # "John Doe at Company LLC"
            r"([A-Z][a-z]+\s+[A-Z][a-z]+)\s+(?:at|from)\s+([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))",  # noqa: E501
            # "John Doe, CEO of Company LLC"
            r"([A-Z][a-z]+\s+[A-Z][a-z]+),?\s+(?:CEO|CFO|CTO|COO|President|Director|Manager|Founder)\s+(?:of|at)\s+([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))",  # noqa: E501
            # "Company LLC - Contact: John Doe"
            r"([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))\s*[-:]\s*(?:Contact|Representative):\s*([A-Z][a-z]+\s+[A-Z][a-z]+)",  # noqa: E501
            # "John Doe (Company LLC)"
            r"([A-Z][a-z]+\s+[A-Z][a-z]+)\s+\(([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))\)",  # noqa: E501
            # Email pattern: john.doe@company.com -> John Doe at Company
            r"([A-Z][a-z]+\s+[A-Z][a-z]+)\s+<[a-zA-Z0-9._%+-]+@([a-zA-Z0-9.-]+)\.[a-zA-Z]{2,}>",
        ]

        for pattern in relationship_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if len(match) == 2:
                    person, company = match
                    # Clean up
                    person_clean = " ".join(person.split())
                    company_clean = " ".join(company.split())

                    # For email domains, capitalize company name
                    if "@" in text and "." in company_clean and len(company_clean.split(".")) >= 2:
                        # This is likely a domain name, extract company name
                        domain_parts = company_clean.split(".")
                        if domain_parts[0].lower() not in [
                            "gmail",
                            "yahoo",
                            "hotmail",
                            "outlook",
                            "mail",
                        ]:
                            company_clean = domain_parts[0].capitalize()

                    # Store relationship (person -> company)
                    if len(person_clean) > 2 and len(company_clean) > 2:
                        relationships[person_clean] = company_clean

        return relationships

    def is_valid_company_name(self, name: str) -> bool:
        """
        Check if a string is a valid company name (not a sentence fragment).

        Returns:
            True if valid company name, False if likely a sentence fragment
        """
        if not name:
            return False

        name_lower = name.lower().strip()
        words = name.split()

        # Reject if too long (real company names are usually < 60 chars)
        if len(name) > 60:
            return False

        # Reject if too many words (company names rarely have > 6 words)
        if len(words) > 6:
            return False

        # Sentence fragment indicators - words that start sentences, not companies
        sentence_starters = {
            "neither",
            "either",
            "total",
            "the",
            "a",
            "an",
            "if",
            "when",
            "where",
            "while",
            "although",
            "because",
            "since",
            "unless",
            "however",
            "therefore",
            "moreover",
            "furthermore",
            "additionally",
            "please",
            "note",
            "see",
            "refer",
            "click",
            "visit",
            "contact",
            "for",
            "with",
            "from",
            "into",
            "about",
            "above",
            "below",
            "between",
            "under",
            "over",
            "after",
            "before",
            "during",
            "this",
            "that",
            "these",
            "those",
            "which",
            "what",
            "who",
            "all",
            "any",
            "each",
            "every",
            "both",
            "few",
            "many",
            "most",
            "other",
            "some",
            "such",
            "no",
            "not",
            "only",
            "own",
            "same",
            "output",
            "input",
            "return",
            "returns",
            "required",
            "optional",
        }

        # Check first word
        if words and words[0].lower() in sentence_starters:
            return False

        # Sentence patterns - these indicate full sentences, not company names
        sentence_patterns = [
            r"\b(?:is|are|was|were|be|been|being)\b",  # Verbs
            r"\b(?:to|of|in|on|at|by)\s+(?:the|a|an)\b",  # Preposition + article
            r"\b(?:you|your|we|our|they|their|it|its)\b",  # Pronouns
            r"\b(?:can|could|will|would|shall|should|may|might|must)\b",  # Modal verbs
            r"\b(?:and|or|but|nor|yet|so)\s+\w+\s+\w+",  # Conjunction + multiple words
        ]

        for pattern in sentence_patterns:
            if re.search(pattern, name_lower):
                return False

        # Check for specific problematic patterns
        problematic_phrases = [
            "the name of",
            "in usd",
            "total in",
            "output only",
            "required for",
            "agreement between",
            "agreement of",
            "certificate of",
            "description of",
            "operating agreement",
            "license this",
            "http rule",
            "member-managed",
            "need some",
            "print out",
            "user provided",
            "ceo of",
            "cfo of",
            "cto of",
            "coo of",  # Title patterns
            "president of",
            "director of",
            "manager of",
            "taxpayer number",
            "tax id",
            "ein number",  # Tax/ID patterns
            "student award",
            "professional access",  # Award patterns
            "proprietor general",
            "general partnership",  # Legal entity types
            "personal workload",
            "workload and",  # Incomplete phrases
            "data usage agreement",
            "service agreement",
            "contributions on behalf",
            "on behalf of",
        ]

        for phrase in problematic_phrases:
            if phrase in name_lower:
                return False

        # Reject names ending with conjunctions (incomplete phrases)
        if words and words[-1].lower() in {
            "and",
            "or",
            "but",
            "nor",
            "yet",
            "so",
            "the",
            "a",
            "an",
            "of",
            "to",
            "in",
            "on",
            "at",
            "by",
        }:
            return False

        # Reject names starting with titles followed by "of"
        if len(words) >= 3 and words[1].lower() == "of":
            title_words = {
                "ceo",
                "cfo",
                "cto",
                "coo",
                "president",
                "director",
                "manager",
                "chairman",
                "founder",
            }
            if words[0].lower() in title_words:
                return False

        return True

    def normalize_company_name(self, company_name: str) -> str:
        """
        Normalize company name by extracting actual company from common patterns.

        Handles patterns like:
        - "Copyright 2024 Google" -> "Google"
        - "© 2020 Microsoft Corporation" -> "Microsoft"
        - "(c) 2019-2024 Apple Inc" -> "Apple"
        - "Copyright (C) 2023 Amazon" -> "Amazon"
        - "Google LLC" -> "Google"
        - "Apple Inc." -> "Apple"

        Returns:
            Normalized company name, or None if invalid
        """
        if not company_name:
            return company_name

        # Patterns to extract company name from copyright notices
        copyright_patterns = [
            # "Copyright 2024 Google" or "Copyright (C) 2024 Google"
            r"(?:copyright|©|\(c\))\s*(?:\(c\))?\s*(?:\d{4}(?:\s*[-–—]\s*\d{4})?)\s+(.+)",
            # "2024 Google" (just year followed by company)
            r"^\d{4}(?:\s*[-–—]\s*\d{4})?\s+([A-Z][A-Za-z0-9\s&\-\.]+)$",
            # "(c) Google 2024" (company before year)
            r"(?:copyright|©|\(c\))\s+([A-Z][A-Za-z0-9\s&\-\.]+?)\s+\d{4}",
            # "Copyright Google" or "© Google" (without year)
            r"^(?:copyright|©|\(c\))\s+([A-Za-z][A-Za-z0-9\s&\-\.]+)$",
        ]

        name_lower = company_name.lower().strip()
        result = company_name

        # Check if this looks like a copyright notice
        if any(indicator in name_lower for indicator in ["copyright", "©", "(c)"]):
            for pattern in copyright_patterns:
                match = re.search(pattern, company_name, re.IGNORECASE)
                if match:
                    extracted = match.group(1).strip()
                    # Clean up trailing punctuation
                    extracted = re.sub(r"[.,;:]+$", "", extracted).strip()
                    if extracted and len(extracted) >= 2:
                        result = extracted
                        break

        # Check for year prefix pattern (e.g., "2024 Google")
        if result == company_name:
            year_prefix_match = re.match(r"^(\d{4}(?:\s*[-–—]\s*\d{4})?)\s+(.+)$", company_name)
            if year_prefix_match:
                extracted = year_prefix_match.group(2).strip()
                if extracted and len(extracted) >= 2:
                    result = extracted

        # Strip legal suffixes to consolidate company variants
        # Order matters: check longer suffixes first
        legal_suffixes = [
            # Full words with variations
            r"\s+Incorporated$",
            r"\s+Corporation$",
            r"\s+Limited$",
            r"\s+Company$",
            # Abbreviations with optional period
            r"\s+L\.L\.C\.$",
            r"\s+L\.L\.P\.$",
            r"\s+LLC\.?$",
            r"\s+LLP\.?$",
            r"\s+Inc\.?$",
            r"\s+Corp\.?$",
            r"\s+Ltd\.?$",
            r"\s+Co\.?$",
            # Other common suffixes
            r"\s+PLC\.?$",
            r"\s+LP\.?$",
            r"\s+SA$",
            r"\s+GmbH$",
            r"\s+AG$",
        ]

        for suffix_pattern in legal_suffixes:
            result = re.sub(suffix_pattern, "", result, flags=re.IGNORECASE).strip()

        return result

    def sanitize_company_name(self, company_name: str) -> Optional[str]:
        """
        Sanitize company name for use in folder names.

        Returns:
            Sanitized folder name, or None if the name is invalid (sentence fragment)
        """
        # First normalize the company name (extract from copyright patterns, etc.)
        normalized = self.normalize_company_name(company_name)

        # Validate that this is a real company name, not a sentence fragment
        if not self.is_valid_company_name(normalized):
            return None

        # Remove special characters that aren't allowed in folder names
        sanitized = re.sub(r'[<>:"/\\|?*]', "", normalized)
        # Replace multiple spaces with single space
        sanitized = " ".join(sanitized.split())
        # Limit length
        if len(sanitized) > 50:
            sanitized = sanitized[:50].strip()
        return sanitized if sanitized else None

    def classify_content(
        self, text: str, filename: str = ""
    ) -> Tuple[str, str, Optional[str], List[str]]:
        """
        Classify content based on extracted text.
        Uses Schema.org person-company relationships to improve categorization.

        Returns:
            Tuple of (category, subcategory, company_name, people_names)
        """
        if not text:
            return ("uncategorized", "other", None, [])

        text_lower = text.lower()
        filename_lower = filename.lower()
        combined = f"{text_lower} {filename_lower}"

        # Check for known companies in text (canonical name mapping)
        known_text_companies = {
            "capital city village": ("organization", "property_management", "Capital City Village"),
            "leora home health": ("organization", "healthcare", "Leora Home Health"),
            "integrity studio": ("organization", "vendors", "Integrity Studio"),
            "inspired movement": ("organization", "vendors", "Inspired Movement"),
            "new beginnings child development": (
                "organization",
                "vendors",
                "New Beginnings Child Development Center",
            ),
            "zouk": ("zouk", "events", None),
        }
        for phrase, (cat, subcat, canonical_name) in known_text_companies.items():
            if phrase in text_lower:
                return (cat, subcat, canonical_name, self.extract_people_names(text))

        # Extract company names and people names
        company_names = self.extract_company_names(text)
        primary_company = company_names[0] if company_names else None

        people_names = self.extract_people_names(text)

        # Extract person-company relationships (Schema.org connections)
        person_company_relationships = self.extract_person_company_relationships(text)

        # Prioritize company from person-company relationships over direct extraction
        # Relationships tend to be more accurate as they include context
        if person_company_relationships:
            # Get the first relationship's company
            relationship_company = next(iter(person_company_relationships.values()))

            # Check if relationship company has proper legal suffix
            has_legal_suffix = any(
                relationship_company.endswith(suffix)
                for suffix in [
                    "LLC",
                    "Inc.",
                    "Inc",
                    "Corp.",
                    "Corp",
                    "Ltd.",
                    "Ltd",
                    "LLP",
                    "L.L.C.",
                    "L.L.P.",
                ]
            )

            # Prefer relationship company if it has legal suffix or we don't have a primary company
            if has_legal_suffix or not primary_company:
                primary_company = relationship_company
            # Or if the relationship company is much cleaner (shorter and no weird prefixes)
            elif (
                primary_company
                and "CEO" not in relationship_company
                and "at" not in relationship_company
            ):
                if len(relationship_company) < len(primary_company) * 0.8:
                    primary_company = relationship_company

        # Fallback: If we found people but no company, check relationships again
        if people_names and not primary_company and person_company_relationships:
            # Try to find a company for the first person mentioned
            for person in people_names:
                if person in person_company_relationships:
                    primary_company = person_company_relationships[person]
                    break

        # Score each category
        scores = defaultdict(int)
        category_subcats = {}

        for category, data in self.patterns.items():
            for keyword in data["keywords"]:
                count = combined.count(keyword.lower())
                if count > 0:
                    scores[category] += count

                    # Track which subcategory keywords matched
                    for subcat, subcat_keywords in data["subcategories"].items():
                        if any(sk.lower() in combined for sk in subcat_keywords):
                            if category not in category_subcats:
                                category_subcats[category] = defaultdict(int)
                            category_subcats[category][subcat] += count

        if not scores:
            return ("uncategorized", "other", primary_company, people_names)

        # Get category with highest score
        best_category = max(scores.items(), key=lambda x: x[1])[0]

        # Get subcategory with highest score for this category
        if best_category in category_subcats:
            subcat_scores = category_subcats[best_category]
            if subcat_scores:
                best_subcategory = max(subcat_scores.items(), key=lambda x: x[1])[0]
            else:
                best_subcategory = "other"
        else:
            best_subcategory = "other"

        # If we detected a company (either directly or via person relationship)
        # and it's business-related, use clients subcategory
        if primary_company and best_category == "business":
            best_subcategory = "clients"

        return (best_category, best_subcategory, primary_company, people_names)


class ImageMetadataParser:
    """Parses image metadata including EXIF, GPS, and timestamps."""

    def __init__(self, cost_calculator: "CostROICalculator" = None):
        """
        Initialize the metadata parser.

        Args:
            cost_calculator: Optional cost calculator for tracking usage costs
        """
        self.metadata_available = METADATA_AVAILABLE
        self.geocoder = None
        self.cost_calculator = cost_calculator

        if self.metadata_available:
            try:
                # Initialize geocoder with a user agent
                self.geocoder = Nominatim(user_agent="file_organizer_v1.0", timeout=5)
            except Exception as e:
                print(f"Warning: Could not initialize geocoder: {e}")
                self.geocoder = None

    def extract_exif_data(self, image_path: Path) -> Dict[str, Any]:
        """
        Extract EXIF data from an image.

        Returns:
            Dictionary with EXIF data
        """
        if not self.metadata_available:
            return {}

        try:
            image = Image.open(image_path)
            exif_data = {}

            # Get EXIF data
            exif = image._getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    exif_data[tag] = value

            return exif_data

        except Exception as e:
            print(f"  EXIF extraction error: {e}")
            return {}

    def extract_datetime(
        self, image_path: Path, exif_data: Optional[Dict[str, Any]] = None
    ) -> Optional[datetime]:
        """
        Extract the datetime when the photo was taken.

        Args:
            image_path: Path to the image file
            exif_data: Pre-extracted EXIF data to avoid redundant Image.open() calls

        Returns:
            datetime object or None
        """
        if exif_data is None:
            exif_data = self.extract_exif_data(image_path)

        if not exif_data:
            return None

        # Try different datetime tags
        datetime_tags = ["DateTimeOriginal", "DateTimeDigitized", "DateTime"]

        for tag in datetime_tags:
            if tag in exif_data:
                try:
                    # Parse datetime string (format: "2023:11:26 14:30:00")
                    dt_str = str(exif_data[tag])
                    dt = datetime.strptime(dt_str, "%Y:%m:%d %H:%M:%S")
                    return dt
                except (ValueError, TypeError):
                    continue

        return None

    def extract_gps_coordinates(
        self, image_path: Path, exif_data: Optional[Dict[str, Any]] = None
    ) -> Optional[Tuple[float, float]]:
        """
        Extract GPS coordinates from image EXIF data.

        Args:
            image_path: Path to the image file
            exif_data: Pre-extracted EXIF data to avoid redundant Image.open() calls

        Returns:
            Tuple of (latitude, longitude) or None
        """
        if not self.metadata_available:
            return None

        try:
            if exif_data is None:
                exif_data = self.extract_exif_data(image_path)

            if not exif_data:
                return None

            # Get GPS info from pre-extracted EXIF data
            gps_info = {}
            gps_raw = exif_data.get("GPSInfo")
            if gps_raw and isinstance(gps_raw, dict):
                for gps_tag_id, value in gps_raw.items():
                    gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                    gps_info[gps_tag] = value

            if not gps_info:
                return None

            # Convert to decimal degrees
            lat = self._convert_to_degrees(gps_info.get("GPSLatitude"))
            lon = self._convert_to_degrees(gps_info.get("GPSLongitude"))

            if lat is None or lon is None:
                return None

            # Adjust for hemisphere
            if gps_info.get("GPSLatitudeRef") == "S":
                lat = -lat
            if gps_info.get("GPSLongitudeRef") == "W":
                lon = -lon

            return (lat, lon)

        except Exception as e:
            print(f"  GPS extraction error: {e}")
            return None

    def _convert_to_degrees(self, value) -> Optional[float]:
        """
        Convert GPS coordinates to degrees.

        Args:
            value: GPS coordinate in format ((deg, 1), (min, 1), (sec, 1))

        Returns:
            Decimal degrees or None
        """
        if not value:
            return None

        try:
            d = float(value[0][0]) / float(value[0][1])
            m = float(value[1][0]) / float(value[1][1])
            s = float(value[2][0]) / float(value[2][1])

            return d + (m / 60.0) + (s / 3600.0)
        except (IndexError, TypeError, ZeroDivisionError):
            return None

    def get_location_name(self, coordinates: Tuple[float, float]) -> Optional[str]:
        """
        Get location name from GPS coordinates using reverse geocoding.

        Args:
            coordinates: Tuple of (latitude, longitude)

        Returns:
            Location name (city, state, country) or None
        """
        if not self.geocoder:
            return None

        with (
            CostTracker(self.cost_calculator, "nominatim_geocoding")
            if self.cost_calculator
            else nullcontext()
        ):
            try:
                lat, lon = coordinates
                location = self.geocoder.reverse(f"{lat}, {lon}", exactly_one=True)

                if location and location.raw.get("address"):
                    address = location.raw["address"]

                    # Try to get city, state, country
                    parts = []

                    # City
                    city = address.get("city") or address.get("town") or address.get("village")
                    if city:
                        parts.append(city)

                    # State/Region
                    state = address.get("state") or address.get("region")
                    if state:
                        parts.append(state)

                    # Country
                    country = address.get("country")
                    if country:
                        parts.append(country)

                    if parts:
                        return ", ".join(parts)

            except (GeocoderTimedOut, GeocoderServiceError) as e:
                print(f"  Geocoding error: {e}")
            except Exception as e:
                print(f"  Location lookup error: {e}")

            return None

    def get_metadata_summary(self, image_path: Path) -> Dict[str, Any]:
        """
        Get a summary of image metadata.

        Returns:
            Dictionary with datetime, GPS coordinates, and location
        """
        summary = {
            "datetime": None,
            "gps_coordinates": None,
            "location_name": None,
            "year": None,
            "month": None,
            "date_str": None,
        }

        # Extract EXIF once, pass to both datetime and GPS extraction
        exif_data = self.extract_exif_data(image_path)

        # Extract datetime
        dt = self.extract_datetime(image_path, exif_data=exif_data)
        if dt:
            summary["datetime"] = dt
            summary["year"] = dt.year
            summary["month"] = dt.month
            summary["date_str"] = dt.strftime("%Y-%m")

        # Extract GPS
        coords = self.extract_gps_coordinates(image_path, exif_data=exif_data)
        if coords:
            summary["gps_coordinates"] = coords
            # Get location name (with rate limiting consideration)
            location = self.get_location_name(coords)
            if location:
                summary["location_name"] = location

        return summary


class ContentBasedFileOrganizer:
    """Organize files based on content analysis using OCR."""

    def __init__(
        self,
        base_path: str = None,
        organize_by_date: bool = False,
        organize_by_location: bool = False,
        enable_cost_tracking: bool = True,
        db_path: str = "results/file_organization.db",
    ):
        """
        Initialize the organizer.

        Args:
            base_path: Base path for organized files
            organize_by_date: If True, organize photos by date (Photos/2023/11/)
            organize_by_location: If True, organize photos by location when GPS data available
            enable_cost_tracking: If True, track costs and ROI for all features
            db_path: Path to SQLite database for persistent storage
        """
        self.base_path = Path(base_path or "~/Documents").expanduser()

        # Initialize cost tracking if available and enabled
        self.cost_calculator = None
        if enable_cost_tracking and COST_TRACKING_AVAILABLE:
            self.cost_calculator = CostROICalculator()
            print("✓ Cost tracking enabled")

        # Initialize graph store for persistent storage with canonical IDs
        self.graph_store = None
        if GRAPH_STORE_AVAILABLE and db_path:
            self.graph_store = GraphStore(db_path=db_path)
            print(f"✓ Graph store enabled ({db_path})")

        self.enricher = MetadataEnricher()
        self.validator = SchemaValidator()
        self.registry = SchemaRegistry()
        self.classifier = ContentClassifier()
        self.rename_analyzer = ImageAnalyzer(PHOTO_PROFILE)
        self.image_analyzer = ImageContentAnalyzer(cost_calculator=self.cost_calculator)
        self.metadata_parser = ImageMetadataParser(cost_calculator=self.cost_calculator)
        self.stats = defaultdict(int)
        self.ocr_available = OCR_AVAILABLE
        self.organize_by_date = organize_by_date
        self.organize_by_location = organize_by_location
        # Temporary OCR metadata for the current file being processed.
        # Set inside detect_file_category; consumed by _persist_to_graph_store.
        self._last_file_ocr_confidence: Optional[float] = None
        self._last_file_detected_language: Optional[str] = None
        self._last_file_ocr_text: Optional[str] = None  # cached OCR text to avoid re-running
        # Structured per-file outputs (KIE result, research-paper metadata).
        # Keyed dict so new metadata types can land here without growing the
        # attribute surface.
        self._last_file_state: Dict[str, Any] = {}
        # Per-file CLIP enhance cache: (file_path, tuple(labels)) -> List of results
        self._clip_enhance_cache: Dict[tuple, list] = {}

        # Filepath-based classification (checked FIRST before content analysis)
        self.filepath_patterns = {
            # Log files
            ".log": "Technical/Logs",
            ".log.gz": "Technical/Logs",
            ".out": "Technical/Logs",
            # Python
            ".py": "Technical/Python",
            ".pyc": "Technical/Python/Compiled",
            ".pyw": "Technical/Python",
            ".pyx": "Technical/Python",
            ".pyd": "Technical/Python",
            # JavaScript/TypeScript
            ".js": "Technical/JavaScript",
            ".jsx": "Technical/JavaScript",
            ".mjs": "Technical/JavaScript",
            ".cjs": "Technical/JavaScript",
            ".ts": "Technical/TypeScript",
            ".tsx": "Technical/TypeScript",
            # Web
            ".html": "Technical/Web",
            ".htm": "Technical/Web",
            ".css": "Technical/Web",
            ".scss": "Technical/Web",
            ".sass": "Technical/Web",
            ".less": "Technical/Web",
            # Shell scripts
            ".sh": "Technical/Shell",
            ".bash": "Technical/Shell",
            ".zsh": "Technical/Shell",
            ".fish": "Technical/Shell",
            # Config files
            ".json": "Technical/Config",
            ".yaml": "Technical/Config",
            ".yml": "Technical/Config",
            ".toml": "Technical/Config",
            ".ini": "Technical/Config",
            ".conf": "Technical/Config",
            ".config": "Technical/Config",
            ".env": "Technical/Config",
            # Database
            ".sql": "Technical/Database",
            ".db": "Technical/Database",
            ".sqlite": "Technical/Database",
            ".sqlite3": "Technical/Database",
            # Java/Kotlin
            ".java": "Technical/Java",
            ".class": "Technical/Java/Compiled",
            ".jar": "Technical/Java/Archives",
            ".kt": "Technical/Kotlin",
            ".kts": "Technical/Kotlin",
            # C/C++
            ".c": "Technical/C",
            ".cpp": "Technical/C++",
            ".cc": "Technical/C++",
            ".cxx": "Technical/C++",
            ".h": "Technical/C/Headers",
            ".hpp": "Technical/C++/Headers",
            # Go
            ".go": "Technical/Go",
            # Rust
            ".rs": "Technical/Rust",
            # Ruby
            ".rb": "Technical/Ruby",
            ".rake": "Technical/Ruby",
            # PHP
            ".php": "Technical/PHP",
            # Swift
            ".swift": "Technical/Swift",
            # Markdown and docs
            ".md": "Technical/Documentation",
            ".markdown": "Technical/Documentation",
            ".rst": "Technical/Documentation",
            ".adoc": "Technical/Documentation",
            # Version control
            ".gitignore": "Technical/VersionControl",
            ".gitattributes": "Technical/VersionControl",
            # Build/Package files
            "Makefile": "Technical/Build",
            "Dockerfile": "Technical/Build",
            "docker-compose.yml": "Technical/Build",
            "package.json": "Technical/Build",
            "package-lock.json": "Technical/Build",
            "yarn.lock": "Technical/Build",
            "Cargo.toml": "Technical/Build",
            "go.mod": "Technical/Build",
            "requirements.txt": "Technical/Build",
            "Pipfile": "Technical/Build",
            "pyproject.toml": "Technical/Build",
        }

        # Content-based organization structure
        self.category_paths = {
            "legal": {
                "contracts": "Legal/Contracts",
                "real_estate": "Legal/RealEstate",
                "corporate": "Legal/Corporate",
                "other": "Legal/Other",
            },
            "financial": {
                "tax": "Financial/Tax",
                "invoices": "Financial/Invoices",
                "statements": "Financial/Statements",
                "other": "Financial/Other",
            },
            "business": {
                "planning": "Business/Planning",
                "marketing": "Business/Marketing",
                "proposals": "Business/Proposals",
                "presentations": "Business/Presentations",
                "crm": "Business/CRM",
                "hr": "Business/HR",
                "meeting_notes": "Business/MeetingNotes",
                "clients": "Business/Clients",  # Legacy - prefer crm/hr
                "other": "Business/Other",
            },
            "personal": {
                "employment": "Personal/Employment",
                "identification": "Personal/Identification",
                "certificates": "Personal/Certificates",
                "other": "Personal/Other",
            },
            "medical": {
                "records": "Medical/Records",
                "insurance": "Medical/Insurance",
                "prescriptions": "Medical/Prescriptions",
                "other": "Medical/Other",
            },
            "property": {
                "leases": "Property/Leases",
                "maintenance": "Property/Maintenance",
                "other": "Property/Other",
            },
            "education": {
                "coursework": "Education/Coursework",
                "research": "Education/Research",
                "records": "Education/Records",
                "other": "Education/Other",
            },
            RESEARCH_CATEGORY: {
                "arxiv": "Research/Papers/arXiv",
                "ssrn": "Research/Papers/SSRN",
                "doi": "Research/Papers/DOI",
                "other": "Research/Papers/Other",
            },
            "technical": {
                "documentation": "Technical/Documentation",
                "architecture": "Technical/Architecture",
                "config": "Technical/Config",
                "data": "Technical/Data",
                "logs": "Technical/Logs",
                "web": "Technical/Web",
                "software_packages": "Technical/Software_Packages",
                "other": "Technical/Other",
            },
            "creative": {
                "design": "Creative/Design",
                "branding": "Creative/Branding",
                "photos": "Creative/Photos",
                "other": "Creative/Other",
            },
            "property_management": "Property_Management",
            "zouk": {"events": "Zouk/Events", "classes": "Zouk/Classes", "other": "Zouk/Other"},
            # Organization: root folder with entity-named subfolders
            # Structure: Organization/{OrgName}/ for most types
            # Exception: Organization/Clients/{OrgName}/ for clients (nested)
            "organization": {
                "clients": "Organization/Clients",  # Gets nested subfolders
                "vendors": "Organization",  # Root folder, entity name added dynamically
                "partners": "Organization",
                "employers": "Organization",
                "government": "Organization",
                "healthcare": "Organization",
                "property_management": "Organization",
                "financial": "Organization",
                "educational": "Organization",
                "nonprofit": "Organization",
                "meeting_notes": "Organization",  # Gets Meeting Notes subfolder after company
                "other": "Organization",
            },
            # Person: root folder with person-named subfolders
            # Structure: Person/{PersonName}/ for all types
            "person": {
                "contacts": "Person",  # Root folder, person name added dynamically
                "employees": "Person",
                "clients": "Person",
                "family": "Person",
                "references": "Person",
                "travel": "Person/Travel",
                "events": "Person/Events",
                "journal": "Person/Journal",  # Personal writing, dreams, reflections
                "other": "Person",
            },
            "game_assets": {
                "audio": "GameAssets/Audio",
                "music": "GameAssets/Music",
                "sprites": "GameAssets/Sprites",
                "textures": "GameAssets/Textures",
                "fonts": "GameAssets/Fonts",
                "other": "GameAssets/Other",
            },
            "fonts": {
                "truetype": "CreativeWork/Fonts/TrueType",
                "opentype": "CreativeWork/Fonts/OpenType",
                "web": "CreativeWork/Fonts/Web",
                "other": "CreativeWork/Fonts/Other",
            },
            "media": {
                "photos": {
                    "screenshots": {
                        "browser": "Media/Photos/Screenshots/Browser",
                        "terminal": "Media/Photos/Screenshots/Terminal",
                        "code": "Media/Photos/Screenshots/CodeEditors",
                        "docs": "Media/Photos/Screenshots/Docs",
                        "settings": "Media/Photos/Screenshots/Settings",
                        "products": "Media/Photos/Screenshots/Products",
                        "dashboard": "Media/Photos/Screenshots/Dashboards",
                        "chat": "Media/Photos/Screenshots/Chat",
                        "other": "Media/Photos/Screenshots",
                    },
                    "travel": "Media/Photos/Travel",
                    "portraits": "Media/Photos/Portraits",
                    "events": "Media/Photos/Events",
                    "documents": "Media/Photos/Documents",
                    "social": "Media/Photos/Social",
                    "chatgpt": "Media/Photos/ChatGPT",
                    "facebook": "Media/Photos/Facebook",
                    "logos": "Media/Photos/Logos",
                    "stock": "Media/Photos/Stock",
                    "nature": "Media/Photos/Nature",
                    "lifestyle": "Media/Photos/Lifestyle",
                    "products": "Media/Photos/Products",
                    "other": "Media/Photos/Other",
                },
                "videos": {
                    "recordings": "Media/Videos/Recordings",
                    "exports": "Media/Videos/Exports",
                    "screencasts": "Media/Videos/Screencasts",
                    "other": "Media/Videos/Other",
                },
                "audio": {
                    "recordings": "Media/Audio/Recordings",
                    "music": "Media/Audio/Music",
                    "podcasts": "Media/Audio/Podcasts",
                    "other": "Media/Audio/Other",
                },
                "graphics": {
                    "vector": "Media/Graphics/Vector",
                    "icons": "Media/Graphics/Icons",
                    "other": "Media/Graphics/Other",
                },
                "other": "Media/Other",
            },
            "uncategorized": "Uncategorized",
        }

        # Extend screenshot sub-folders from classifier taxonomies so that
        # content labels map directly to folder paths without a separate lookup.
        _screenshots = self.category_paths["media"]["photos"]["screenshots"]
        for key in SCREENSHOT_KEYWORDS:
            if key not in _screenshots:
                folder = key.replace("_", " ").title().replace(" ", "")
                _screenshots[key] = f"Media/Photos/Screenshots/{folder}"
        for key in self.classifier.patterns:
            if key not in _screenshots:
                _screenshots[key] = f"Media/Photos/Screenshots/{key.title()}"

        # Game asset detection patterns
        self.game_audio_keywords = [
            "bolt",
            "spell",
            "magic",
            "cast",
            "chirp",
            "crossbow",
            "dagger",
            "sword",
            "arrow",
            "bow",
            "heal",
            "potion",
            "lightning",
            "fire",
            "ice",
            "acid",
            "poison",
            "explosion",
            "blast",
            "summon",
            "dispel",
            "petrification",
            "neutralize",
            "slow",
            "darkness",
            "achievement",
            "quest",
            "unlock",
            "lock",
            "door",
            "chest",
            "coin",
            "pickup",
            "attack",
            "hit",
            "damage",
            "death",
            "footstep",
            "jump",
            "land",
            "monster",
            "creature",
            "enemy",
            "boss",
            "battle",
            "combat",
            "starving",
            "hunger",
            "thirst",
            "eat",
            "drink",
            "sleep",
            "fiddle",
            "lute",
            "mandoline",
            "glockenspiel",
            "instrument",
            "identify",
            "greater",
            "mental",
        ]

        self.game_music_keywords = [
            "battle",
            "boss",
            "dungeon",
            "castle",
            "forest",
            "town",
            "village",
            "temple",
            "ruins",
            "cave",
            "mountain",
            "ocean",
            "desert",
            "snow",
            "victory",
            "defeat",
            "theme",
            "menu",
            "credits",
            "intro",
            "outro",
            "mysterious",
            "dark",
            "light",
            "epic",
            "calm",
            "peaceful",
            "tension",
            "chaos",
            "hope",
            "despair",
            "triumph",
            "march",
            "symphony",
            "monotony",
            "drakalor",
            "altar",
            "lawful",
            "chaotic",
            "neutral",
            "alignment",
            "dwarven",
            "elven",
            "orcish",
            "halls",
            "abandon",
            "corrupting",
            "breeze",
            "clockwork",
            "knowledge",
            "oddisey",
            "final",
            "welcome",
        ]

        from shared.constants import GAME_SPRITE_KEYWORDS

        self.game_sprite_keywords = GAME_SPRITE_KEYWORDS

        # Regex patterns for game asset detection (numbered sprites, variants)
        import re

        self.game_sprite_patterns = [
            re.compile(r"^\d+_\d+$"),  # 42_8, 51_3, 16_3 (sprite sheets)
            re.compile(r"^\d+_grey(_\d+)?$", re.IGNORECASE),  # 10_grey, 10_grey_1
            re.compile(r"^\d+_f(_\d+)?$", re.IGNORECASE),  # 283_f, 283_f_1
            re.compile(r"^[a-z]+_\d+$", re.IGNORECASE),  # frame_1, item_42
            re.compile(r"^[a-z]+_[a-z]+_\d+$", re.IGNORECASE),  # assassins_deed_1
            re.compile(r"^\d+h_[a-z]+(_\d+)?$", re.IGNORECASE),  # 2h_axe, 2h_axe_1
            re.compile(r"^[a-z]+_v(_\d+)?$", re.IGNORECASE),  # arrow_v, arrow_v_1
            re.compile(r"^[a-z]+_h(_\d+)?$", re.IGNORECASE),  # arrow_h, arrow_h_1
            re.compile(r"^(head|torso|arm|leg|body|wing|hair)_\w+", re.IGNORECASE),  # body parts
            re.compile(
                r"^(weapon|armor|item|sprite|frame|tile)\d*_", re.IGNORECASE
            ),  # game prefixes
        ]

        # Game font sprite sheet patterns
        self.game_font_keywords = [
            "broguefont",
            "gamefont",
            "pixelfont",
            "bitfont",
            "font_",
            "_font",
            "fontsheet",
            "font_atlas",
            "fontatlas",
            "charset",
            "glyphs",
            "tilefont",
            "asciifont",
            "ascii_font",
        ]

    def classify_by_filepath(self, file_path: Path) -> Optional[str]:
        """
        Classify file based on filepath patterns (extension, filename).

        Returns:
            Category path string if matched, None otherwise
        """
        # Check exact filename matches first (e.g., Makefile, Dockerfile)
        filename = file_path.name
        if filename in self.filepath_patterns:
            return self.filepath_patterns[filename]

        # Check file extension
        ext = file_path.suffix.lower()
        if ext in self.filepath_patterns:
            base_path = self.filepath_patterns[ext]

            # Try to extract project name from path
            project_name = self.extract_project_name(file_path)
            if project_name:
                # Add project subdirectory (e.g., Technical/Python/MyProject)
                return f"{base_path}/{project_name}"

            return base_path

        # Check double extensions (e.g., .log.gz)
        if len(file_path.suffixes) >= 2:
            double_ext = "".join(file_path.suffixes[-2:]).lower()
            if double_ext in self.filepath_patterns:
                return self.filepath_patterns[double_ext]

        return None

    def extract_project_name(self, file_path: Path) -> Optional[str]:
        """
        Extract project name from file path.

        Looks for common project indicators in path:
        - Directory names like 'myproject', 'my-app', etc.
        - Skips common non-project directories

        Returns:
            Project name if found, None otherwise
        """
        skip_dirs = {
            "src",
            "lib",
            "bin",
            "dist",
            "build",
            "out",
            "target",
            "node_modules",
            "venv",
            ".venv",
            "env",
            "__pycache__",
            "scripts",
            "tests",
            "test",
            "docs",
            "doc",
            "examples",
            "static",
            "public",
            "assets",
            "resources",
            "config",
            "home",
            "users",
            "documents",
            "downloads",
            "desktop",
            "code",
            "projects",
            "dev",
            "work",
            "repos",
            "git",
        }

        # Get all parent directories
        parts = file_path.parts

        # Look backwards from the file for a likely project directory
        for i in range(len(parts) - 2, -1, -1):  # Skip the filename itself
            dir_name = parts[i].lower()

            # Skip common non-project directories
            if dir_name in skip_dirs:
                continue

            # Skip hidden directories
            if dir_name.startswith("."):
                continue

            # Found a likely project directory
            # Return with original case preserved
            return parts[i]

        return None

    def classify_game_asset(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Classify file as a game asset based on filename patterns.

        Returns:
            Tuple of (category, subcategory) or None if not a game asset
        """
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()

        # Remove timestamp suffixes for pattern matching (e.g., _20251120_164506)
        import re

        clean_stem = re.sub(r"_\d{8}_\d{6}$", "", stem)

        # Check for audio files (.wav, .ogg, .mp3)
        if ext in [".wav", ".ogg", ".mp3", ".flac", ".aac"]:
            # Check for game music patterns (usually .ogg files with specific names)
            if ext == ".ogg":
                for keyword in self.game_music_keywords:
                    if keyword in stem:
                        return ("game_assets", "music")

            # Check for game sound effects
            for keyword in self.game_audio_keywords:
                if keyword in stem:
                    return ("game_assets", "audio")

        # Check for image files that are game sprites/textures
        # Exclude files with 'screenshot' in name — those are screen captures
        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".tga", ".dds"] and "screenshot" not in stem:
            # Check for game font sprite sheets first
            for keyword in self.game_font_keywords:
                if keyword in stem or keyword in clean_stem:
                    return ("game_assets", "fonts")

            # Check regex patterns for numbered sprites and variants
            for pattern in self.game_sprite_patterns:
                if pattern.match(clean_stem):
                    return ("game_assets", "sprites")

            # Check for sprite/texture keyword patterns
            for keyword in self.game_sprite_keywords:
                if keyword in stem or keyword in clean_stem:
                    # Distinguish between sprites and textures
                    sprite_keywords = [
                        "frame",
                        "sprite",
                        "leg",
                        "arm",
                        "head",
                        "torso",
                        "body",
                        "wing",
                        "hair",
                        "face",
                        "mouth",
                        "_grey",
                        "_gray",
                        "assassins",
                        "atonement",
                        "arrow_v",
                        "arrow_h",
                        "add",
                        "2h_",
                        "1h_",
                        "dagger",
                        "sword",
                        "axe",
                        "hammer",
                        "mace",
                        # Character customization
                        "beard",
                        "bling",
                        "hiero",
                        "mustache",
                        "scar",
                        "tattoo",
                        "earring",
                        "necklace",
                        "bracelet",
                        "glasses",
                        "mask",
                        "hood",
                    ]
                    if any(kw in stem or kw in clean_stem for kw in sprite_keywords):
                        return ("game_assets", "sprites")
                    else:
                        return ("game_assets", "textures")

        # Check for font files
        if ext in [".ttf", ".otf", ".woff", ".woff2", ".eot", ".fon", ".fnt"]:
            if ext == ".ttf":
                return ("fonts", "truetype")
            elif ext == ".otf":
                return ("fonts", "opentype")
            elif ext in [".woff", ".woff2", ".eot"]:
                return ("fonts", "web")
            else:
                return ("fonts", "other")

        return None

    def classify_by_organization(self, text: str, filename: str) -> Optional[Tuple[str, str, str]]:
        """
        Classify file primarily by Organization entity detection.

        Looks for strong organization indicators like:
        - Company names in headers/footers
        - Official letterheads
        - Business correspondence
        - Invoices, contracts with company names

        Returns:
            Tuple of (category, subcategory, org_name) or None if no strong organization match
        """
        if not text or len(text) < 50:
            return None

        text_lower = text.lower()

        # Organization type indicators
        org_indicators = {
            "government": [
                "department of",
                "internal revenue",
                "irs",
                "social security",
                "state of",
                "county of",
                "city of",
                "municipality",
                "federal",
                "government",
                "agency",
                "bureau",
                "commission",
                "dmv",
                "passport",
                "immigration",
                "customs",
                "treasury",
            ],
            "healthcare": [
                "hospital",
                "clinic",
                "medical center",
                "health system",
                "healthcare",
                "physicians",
                "doctor",
                "patient",
                "diagnosis",
                "prescription",
                "pharmacy",
                "insurance claim",
                "medicare",
                "medicaid",
                "hipaa",
                "medical record",
                "lab results",
            ],
            "financial": [
                "bank",
                "credit union",
                "investment",
                "brokerage",
                "mortgage",
                "loan",
                "account statement",
                "transaction",
                "wire transfer",
                "routing number",
                "account number",
                "fdic",
                "securities",
            ],
            "educational": [
                "university",
                "college",
                "school",
                "academy",
                "institute",
                "transcript",
                "diploma",
                "degree",
                "enrollment",
                "registrar",
                "financial aid",
                "tuition",
                "semester",
                "course",
                "student id",
            ],
            "nonprofit": [
                "foundation",
                "charity",
                "nonprofit",
                "non-profit",
                "501(c)",
                "donation",
                "volunteer",
                "mission",
                "charitable",
            ],
            "employers": [
                "offer letter",
                "employment agreement",
                "w-2",
                "w2",
                "pay stub",
                "payroll",
                "human resources",
                "hr department",
                "employee id",
                "benefits enrollment",
                "performance review",
                "termination",
            ],
            "vendors": [
                "invoice",
                "purchase order",
                "po number",
                "vendor id",
                "supplier",
                "bill to",
                "ship to",
                "payment terms",
                "net 30",
            ],
            "clients": [
                "client",
                "customer",
                "service agreement",
                "statement of work",
                "sow",
                "proposal",
                "quote",
                "estimate",
                "engagement letter",
            ],
        }

        # Check for organization type indicators
        for org_type, keywords in org_indicators.items():
            matches = sum(1 for kw in keywords if kw in text_lower)
            if matches >= 2:  # Require at least 2 keyword matches
                # Try to extract organization name
                companies = self.classifier.extract_company_names(text)
                org_name = companies[0] if companies else None
                if org_name:
                    return ("organization", org_type, org_name)

        return None

    def classify_by_person(self, text: str, filename: str) -> Optional[Tuple[str, str, List[str]]]:
        """
        Classify file primarily by Person entity detection.

        Looks for strong person indicators like:
        - Resumes/CVs
        - Contact information (vCards)
        - Personal identification documents
        - Reference letters

        Returns:
            Tuple of (category, subcategory, person_names) or None if no strong person match
        """
        if not text or len(text) < 50:
            return None

        text_lower = text.lower()
        filename_lower = filename.lower()

        # Person type indicators
        person_indicators = {
            "contacts": [
                "contact",
                "phone:",
                "email:",
                "address:",
                "mobile:",
                "tel:",
                "fax:",
                "linkedin",
                "twitter",
                "@",
            ],
            "employees": [
                "employee",
                "staff",
                "team member",
                "department:",
                "title:",
                "hire date",
                "start date",
                "position:",
                "role:",
            ],
            "references": [
                "reference",
                "recommendation",
                "letter of",
                "to whom it may concern",
                "i am pleased to",
                "i highly recommend",
                "worked with",
            ],
            "clients": [
                "client profile",
                "customer profile",
                "client information",
                "account holder",
                "policyholder",
            ],
        }

        # Check filename patterns for resumes/CVs
        resume_patterns = ["resume", "cv", "curriculum", "vitae"]
        if any(pat in filename_lower for pat in resume_patterns):
            people = self.classifier.extract_people_names(text)
            return ("person", "contacts", people if people else [])

        # Check for person type indicators
        for person_type, keywords in person_indicators.items():
            matches = sum(1 for kw in keywords if kw in text_lower)
            if matches >= 2:  # Require at least 2 keyword matches
                people = self.classifier.extract_people_names(text)
                if people and _has_human_name_signal(text):
                    return ("person", person_type, people)

        return None

    def classify_media_file(
        self, file_path: Path, image_metadata: Dict = None
    ) -> Optional[Tuple[str, str, str]]:
        """
        Classify media files (photos, videos, audio) into subcategories.

        Returns:
            Tuple of (category, media_type, subcategory) or None if not a media file
            Example: ('media', 'photos', 'screenshots') or ('media', 'videos', 'recordings')
        """
        filename = file_path.name.lower()
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()

        # Videos - .mp4, .mov, .avi, .mkv, .webm, .m4v
        if ext in [".mp4", ".mov", ".avi", ".mkv", ".webm", ".m4v", ".flv", ".wmv"]:
            # Screen recordings
            if "screen" in stem or "recording" in stem or "capture" in stem:
                return ("media", "videos", "screencasts")
            # Exports (from video editors)
            elif "export" in stem or "render" in stem or "final" in stem or "cut" in stem:
                return ("media", "videos", "exports")
            # Default to recordings
            else:
                return ("media", "videos", "recordings")

        # Audio - .mp3, .wav, .m4a, .aac, .flac, .ogg (but not game music)
        if ext in [".mp3", ".m4a", ".aac", ".flac", ".wma"]:
            # Podcasts
            if "podcast" in stem or "episode" in stem or "interview" in stem:
                return ("media", "audio", "podcasts")
            # Music
            elif "song" in stem or "album" in stem or "track" in stem or "music" in stem:
                return ("media", "audio", "music")
            # Voice recordings
            elif "recording" in stem or "voice" in stem or "memo" in stem or "audio" in stem:
                return ("media", "audio", "recordings")
            # Default to recordings
            else:
                return ("media", "audio", "recordings")

        # Photos - .jpg, .jpeg, .png, .heic, .gif, .webp, .bmp
        if ext in [".jpg", ".jpeg", ".png", ".heic", ".gif", ".webp", ".bmp", ".tiff", ".tif"]:
            # Screenshots — fall through to None so CLIP/OCR sub-classification
            # at Priority 4.5 can route to Browser/Terminal/Docs/etc.
            if (
                filename.startswith("screenshot")
                or "screen shot" in filename
                or "screenshot" in stem
            ):
                return None

            # Scanned documents/receipts (OCR will detect text)
            if "scan" in stem or "receipt" in stem or "document" in stem or "invoice" in stem:
                return ("media", "photos", "documents")

            # Travel photos (has GPS metadata)
            if image_metadata and image_metadata.get("gps_coordinates"):
                # If we have GPS coordinates, it's likely a travel photo
                return ("media", "photos", "travel")

            # Photos with datetime (camera photos) - organize by type
            if image_metadata and image_metadata.get("datetime"):
                # Photos with camera EXIF data are likely personal photos
                # Default to 'other' category for general photos
                return ("media", "photos", "other")

            # Photos without metadata - still categorize as media if they're actual photos
            # (as opposed to game sprites which would be caught earlier)
            if ext in [".jpg", ".jpeg", ".heic"]:
                return ("media", "photos", "other")

            # PNG files without clear classification fall through
            # (could be screenshots, documents, or game assets that weren't caught)
            return None

        return None

    def classify_by_filename_patterns(
        self, file_path: Path
    ) -> Optional[Tuple[str, str, Optional[str], List[str]]]:
        """
        Classify file based on filename patterns before content extraction.

        Delegates to the shared ``filename_classifier`` so the rule set is
        single-homed (see ``src/organizers/content_organizer.py`` for the other
        caller). Returns ``(category, subcategory, company_name, people_names)``
        or ``None``.
        """
        return _classify_by_filename_patterns(
            file_path,
            game_sprite_keywords=self.game_sprite_keywords,
            last_file_state=self._last_file_state,
        )

    def extract_text_from_image(self, image_path: Path) -> str:
        """Extract text from image using docTR OCR.

        Reuses cached OCR text from ID detection or the image renamer
        when available, avoiding a redundant OCR pass.
        """
        if not self.ocr_available:
            return ""

        # Return cached OCR text from an earlier pipeline stage if available.
        if self._last_file_ocr_text:
            return self._last_file_ocr_text

        with (
            CostTracker(self.cost_calculator, "doctr_ocr")
            if self.cost_calculator
            else nullcontext()
        ):
            try:
                result = extract_ocr_text(image_path, max_chars=0)
                return result or ""
            except Exception as e:
                print(f"  OCR error: {e}")
                return ""

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF (searchable or scanned)."""
        if not self.ocr_available:
            return ""

        with (
            CostTracker(self.cost_calculator, "pdf_extraction")
            if self.cost_calculator
            else nullcontext()
        ):
            text = ""

            try:
                # First try to extract text directly (for searchable PDFs)
                with open(pdf_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages[:10]:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                # If we got meaningful text, return it
                if len(text.strip()) > 100:
                    return text.strip()

                # Otherwise, try docTR OCR on the PDF
                print("  Using docTR OCR for scanned PDF...")
                ocr_text = extract_ocr_text_pdf(pdf_path, max_pages=5)
                if ocr_text:
                    text += ocr_text

                return text.strip()
            except Exception as e:
                print(f"  PDF extraction error: {e}")
                return ""

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            return ""

        with (
            CostTracker(self.cost_calculator, "docx_extraction")
            if self.cost_calculator
            else nullcontext()
        ):
            try:
                doc = Document(docx_path)
                text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text.append(paragraph.text)

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text.append(cell.text)

                return "\n".join(text)
            except Exception as e:
                print(f"  DOCX extraction error: {e}")
                return ""

    def extract_text_from_xlsx(self, xlsx_path: Path) -> str:
        """Extract text from Excel spreadsheet."""
        if not EXCEL_AVAILABLE:
            return ""

        with (
            CostTracker(self.cost_calculator, "xlsx_extraction")
            if self.cost_calculator
            else nullcontext()
        ):
            try:
                workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
                text = []

                # Limit to first 5 sheets
                for sheet_name in list(workbook.sheetnames)[:5]:
                    sheet = workbook[sheet_name]
                    # Limit to first 100 rows
                    for row in list(sheet.iter_rows(max_row=100, values_only=True)):
                        row_text = " ".join([str(cell) for cell in row if cell is not None])
                        if row_text.strip():
                            text.append(row_text)

                workbook.close()
                return "\n".join(text)
            except Exception as e:
                print(f"  XLSX extraction error: {e}")
                return ""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from various file types."""
        mime_type = self.enricher.detect_mime_type(str(file_path))
        file_ext = file_path.suffix.lower()

        # Images
        if mime_type and mime_type.startswith("image/"):
            return self.extract_text_from_image(file_path)

        # PDFs
        elif mime_type == "application/pdf" or file_ext == ".pdf":
            return self.extract_text_from_pdf(file_path)

        # Word documents
        elif file_ext in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path)

        # Excel spreadsheets
        elif file_ext in [".xlsx", ".xls"]:
            return self.extract_text_from_xlsx(file_path)

        # Text files
        elif mime_type and mime_type.startswith("text/") or file_ext in [".txt", ".md", ".csv"]:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read(50000)  # First 50KB
            except Exception:
                return ""

        return ""

    _GEOGRAPHIC_LABELS = frozenset(
        {
            "a landscape or nature scene",
            "a cityscape or urban scene",
            "a building or architecture",
        }
    )

    def _map_clip_label(self, label: str, image_metadata: Dict = None) -> Optional[Tuple[str, str]]:
        """Map a CLIP label to (category, subcategory), upgrading to travel if GPS present."""
        mapping = CLIP_LABEL_TO_ORGANIZER.get(label)
        if not mapping:
            return None
        cat, subcat = mapping
        if (
            image_metadata
            and image_metadata.get("gps_coordinates")
            and label in self._GEOGRAPHIC_LABELS
        ):
            cat, subcat = "media", "photos_travel"
        return (cat, subcat)

    def _merge_clip_text_scores(
        self,
        clip_candidate: Optional[Tuple[str, str]],
        clip_score: float,
        text_candidate: Optional[Tuple[str, str]],
        text_chars: int,
    ) -> Optional[Tuple[Tuple[str, str], float, str]]:
        """Combine CLIP and OCR-text signals into a per-(category, subcategory)
        weighted score. Returns (winner, total_score, sources_str) or None.

        Used by both the main image path and enhance_weak_image_classification
        so that PDF/PNG sibling files of the same content converge on the same
        category regardless of which extraction modality dominates.
        """
        scores: Dict[Tuple[str, str], float] = {}
        sources: Dict[Tuple[str, str], List[str]] = {}
        if clip_candidate and clip_score > 0:
            scores[clip_candidate] = scores.get(clip_candidate, 0.0) + clip_score
            sources.setdefault(clip_candidate, []).append(f"CLIP {clip_score:.2f}")
        if text_candidate and text_chars >= _TEXT_MIN_CHARS:
            text_score = _TEXT_SIGNAL_PRIOR * min(1.0, text_chars / _TEXT_LENGTH_FULL_CHARS)
            scores[text_candidate] = scores.get(text_candidate, 0.0) + text_score
            sources.setdefault(text_candidate, []).append(f"text {text_score:.2f}")
        if clip_candidate and text_candidate and clip_candidate == text_candidate:
            scores[clip_candidate] += _SIGNAL_AGREEMENT_BOOST
            sources[clip_candidate].append("agree")
        if not scores:
            return None
        winner, top = max(scores.items(), key=lambda kv: kv[1])
        return winner, top, ", ".join(sources[winner])

    def _cross_check_with_clip(
        self,
        file_path: Path,
        image_metadata: Optional[Dict],
        category: str,
        subcategory: str,
        text_chars: int,
    ) -> Tuple[str, str]:
        """Cross-check a text-derived image classification against CLIP and
        return the merged winner, swapping (cat, sub) only when CLIP outscores
        the text candidate.
        """
        clip_candidate, clip_score = self._run_clip_signal(file_path, image_metadata)
        if not clip_candidate:
            return (category, subcategory)
        merged = self._merge_clip_text_scores(
            clip_candidate,
            clip_score,
            (category, subcategory),
            text_chars,
        )
        if not merged:
            return (category, subcategory)
        (m_cat, m_sub), m_score, m_src = merged
        if (m_cat, m_sub) == (category, subcategory):
            return (category, subcategory)
        print(f"  Unified → {m_cat}/{m_sub} ({m_src}; total {m_score:.2f})")
        return (m_cat, m_sub)

    def _run_clip_signal(
        self,
        file_path: Path,
        image_metadata: Optional[Dict] = None,
    ) -> Tuple[Optional[Tuple[str, str]], float]:
        """Run the 20-category CLIP classifier and map its top label.

        Returns (candidate, score) or (None, 0.0) if CLIP is unavailable or
        below the enhancement threshold.
        """
        if not ENHANCED_CLIP_AVAILABLE or not self.image_analyzer.vision_available:
            return (None, 0.0)
        enhance_cache_key = (str(file_path), tuple(CLIP_CATEGORY_PROMPTS))
        try:
            if enhance_cache_key in self._clip_enhance_cache:
                results = self._clip_enhance_cache[enhance_cache_key]
            elif CLIP_CACHE_AVAILABLE:
                results = get_cached_embedding(file_path, CLIP_CATEGORY_PROMPTS, prompt_prefix="")
                self._clip_enhance_cache[enhance_cache_key] = results
            else:
                results = get_clip_classifier().classify_raw(file_path, CLIP_CATEGORY_PROMPTS)
                self._clip_enhance_cache[enhance_cache_key] = results
            best_prompt, best_score = results[0]
            best_label = _CLIP_PROMPT_TO_LABEL.get(best_prompt, best_prompt)
        except Exception as e:
            print(f"  CLIP signal error: {e}")
            return (None, 0.0)
        if best_score < CLIP_ENHANCE_THRESHOLD:
            return (None, 0.0)
        return (self._map_clip_label(best_label, image_metadata), best_score)

    def enhance_weak_image_classification(
        self, file_path: Path, image_metadata: Dict = None
    ) -> Optional[Tuple[str, str]]:
        """Run full 20-category CLIP + OCR fallback for weakly classified images.

        Only called for images that would otherwise land in photos_other or uncategorized.
        Returns (category, subcategory) or None to keep original classification.
        """
        clip_candidate, clip_score = self._run_clip_signal(file_path, image_metadata)
        if clip_score <= 0:
            return None

        text_candidate: Optional[Tuple[str, str]] = None
        text_chars = 0
        if self.ocr_available:
            try:
                ocr_text = self.extract_text_from_image(file_path)
                if ocr_text and len(ocr_text) >= _TEXT_MIN_CHARS:
                    text_chars = len(ocr_text)
                    text_cat, text_subcat, _, _ = self.classifier.classify_content(
                        ocr_text, file_path.name
                    )
                    if text_cat != "uncategorized":
                        text_candidate = (text_cat, text_subcat)
            except Exception as e:
                print(f"  CLIP enhance OCR error: {e}")

        merged = self._merge_clip_text_scores(
            clip_candidate,
            clip_score,
            text_candidate,
            text_chars,
        )
        if not merged:
            return None
        winner, top_score, src = merged
        print(f"  Unified → {winner[0]}/{winner[1]} ({src}; total {top_score:.2f})")
        return winner

    def detect_file_category(
        self,
        file_path: Path,
        display_path: Path | None = None,
    ) -> Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]:
        """
        Detect file category based on content.

        Args:
            file_path: Physical path to the file on disk (used for content reading).
            display_path: Optional override path whose *name* is used for
                filename-pattern classification.  When the image renamer proposes
                a rename in dry-run mode, display_path carries the descriptive
                name while file_path still points to the original file.

        Priority order (executed in this exact sequence; first match wins):
        0a. Renamed-screenshot content routing (display_path vs file_path)
        0b. Filename pattern detection (fastest - no content extraction needed)
        1.  Organization / Person entity detection (document-type files)
        3.  Game asset detection (audio, sprites, textures)
        3.  Filepath-based classification (file extensions, filenames)
        3.5 Identification-document detection via OCR (passport, ID, license)
        4.  Media file classification (photos, videos, audio)
        4.5 Screenshot sub-classification via OCR + CLIP
        5.  Photo composition analysis (people / home interior)
        6.  Regular text extraction and content/KIE classification

        Returns:
            Tuple of (main_category, subcategory, schema_type, extracted_text,
            company_name, people_names, image_metadata)
        """
        self._last_file_ocr_confidence = None
        self._last_file_detected_language = None
        self._last_file_ocr_text = None
        self._last_file_state.clear()
        self._clip_enhance_cache.clear()

        pattern_path = display_path or file_path

        # Determine schema type and MIME type early (needed for multiple paths)
        mime_type = self.enricher.detect_mime_type(str(file_path))
        if mime_type:
            if mime_type.startswith("image/"):
                schema_type = "ImageObject"
            elif mime_type == "application/pdf":
                schema_type = "DigitalDocument"
            elif mime_type.startswith("video/"):
                schema_type = "VideoObject"
            elif mime_type.startswith("audio/"):
                schema_type = "AudioObject"
            else:
                schema_type = "DigitalDocument"
        else:
            schema_type = "DigitalDocument"

        # PRIORITY 0a: Renamed screenshots → route to category sub-folder
        # When the image renamer classified a screenshot (e.g. "Screenshot ..." →
        # "20260320_terminal_session.png"), match the content label against
        # _SCREENSHOT_KEYWORDS and ContentClassifier.patterns keys directly.
        if display_path and display_path != file_path and "screenshot" in file_path.stem.lower():
            renamed_stem = display_path.stem.lower()
            screenshots_dict = self.category_paths["media"]["photos"]["screenshots"]
            # Check longer keys first so "terminal_session" matches before "terminal"
            for key in sorted(screenshots_dict, key=len, reverse=True):
                if key != "other" and key in renamed_stem:
                    print(f"  ✓ Screenshot content: {key}")
                    return ("media", f"photos_screenshots_{key}", schema_type, "", None, [], {})

        # PRIORITY 0b: Filename pattern detection (fastest - no content extraction needed)
        # Handles: Google invoices, resumes, technical files, legal docs,
        # business docs, entity files
        filename_result = self.classify_by_filename_patterns(pattern_path)
        if filename_result:
            category, subcategory, company_name, people_names = filename_result
            # Handle skip category for duplicates
            if category == "skip":
                return ("skip", subcategory, schema_type, "", None, [], {})
            if category == RESEARCH_CATEGORY:
                schema_type = SCHOLARLY_ARTICLE_SCHEMA_TYPE
            # Point A: enhance weak photos_other from filename patterns for images
            if subcategory == "photos_other" and schema_type == "ImageObject":
                enhanced = self.enhance_weak_image_classification(file_path)
                if enhanced:
                    return (enhanced[0], enhanced[1], schema_type, "", None, [], {})
            return (category, subcategory, schema_type, "", company_name, people_names, {})

        # PRIORITY 1: Organization and Person detection for document-type files
        # Only apply to document/PDF types (not images, audio, video)
        if schema_type == "DigitalDocument" or mime_type == "application/pdf":
            print("  Checking for Organization/Person entities...")
            extracted_text = self.extract_text(file_path)

            if extracted_text and len(extracted_text) >= 50:
                # Try Organization detection first
                org_result = self.classify_by_organization(extracted_text, file_path.name)
                if org_result:
                    category, subcategory, org_name = org_result
                    print(f"  ✓ Organization detected: {org_name} ({subcategory})")
                    return (category, subcategory, schema_type, extracted_text, org_name, [], {})

                # Try Person detection second
                person_result = self.classify_by_person(extracted_text, file_path.name)
                if person_result:
                    category, subcategory, people_names = person_result
                    print(
                        f"  ✓ Person detected: {', '.join(people_names[:3]) if people_names else 'Unknown'} ({subcategory})"  # noqa: E501
                    )
                    return (
                        category,
                        subcategory,
                        schema_type,
                        extracted_text,
                        None,
                        people_names,
                        {},
                    )

        # PRIORITY 3: Check for game assets (before filepath patterns)
        game_asset = self.classify_game_asset(file_path)
        if game_asset:
            category, subcategory = game_asset
            print(f"  ✓ Game asset detected: {subcategory}")
            return (category, subcategory, schema_type, "", None, [], {})

        # PRIORITY 3: Check filepath patterns (most efficient and accurate for code files)
        filepath_category = self.classify_by_filepath(file_path)
        if filepath_category:
            print(f"  ✓ Filepath match: {filepath_category}")
            # Return filepath-based category as a special marker
            # We'll handle this in get_destination_path
            return ("filepath", filepath_category, schema_type, "", None, [], {})

        # Extract metadata for images
        image_metadata = {}
        if schema_type == "ImageObject" and self.metadata_parser.metadata_available:
            print("  Extracting image metadata...")
            image_metadata = self.metadata_parser.get_metadata_summary(file_path)

            if image_metadata.get("datetime"):
                dt = image_metadata["datetime"]
                print(f"  ✓ Photo taken: {dt.strftime('%Y-%m-%d %H:%M:%S')}")

            if image_metadata.get("gps_coordinates"):
                coords = image_metadata["gps_coordinates"]
                print(f"  ✓ GPS: {coords[0]:.6f}, {coords[1]:.6f}")

            if image_metadata.get("location_name"):
                print(f"  ✓ Location: {image_metadata['location_name']}")

        # PRIORITY 3.5: Check for identification documents in images (passport, ID, license)
        # These should go to Person/ folder, not Media/
        id_result = self._classify_identification_document(
            file_path,
            schema_type,
            image_metadata,
        )
        if id_result is not None:
            return id_result

        # PRIORITY 4: Check for media files (photos, videos, audio)
        # This runs after metadata extraction so we can use GPS/datetime for classification
        media_classification = self.classify_media_file(file_path, image_metadata)
        if media_classification:
            category, media_type, subcategory = media_classification
            # Point B: enhance weak photos/other for images
            if media_type == "photos" and subcategory == "other":
                enhanced = self.enhance_weak_image_classification(file_path, image_metadata)
                if enhanced:
                    print(f"  ✓ Enhanced media: {enhanced[0]}/{enhanced[1]}")
                    return (enhanced[0], enhanced[1], schema_type, "", None, [], image_metadata)
            print(f"  ✓ Media file detected: {media_type}/{subcategory}")
            return (
                category,
                f"{media_type}_{subcategory}",
                schema_type,
                "",
                None,
                [],
                image_metadata,
            )

        # PRIORITY 4.5: Screenshot sub-classification via OCR + CLIP
        screenshot_result = self._classify_screenshot_ocr(
            file_path,
            schema_type,
            image_metadata,
        )
        if screenshot_result is not None:
            return screenshot_result

        # PRIORITY 5: Check for photos with people (social) / home interior
        photo_result = self._classify_photo_composition(
            file_path,
            schema_type,
            image_metadata,
        )
        if photo_result is not None:
            return photo_result

        # PRIORITY 6: Regular text extraction and classification
        return self._classify_by_content_and_kie(
            file_path,
            schema_type,
            image_metadata,
        )

    def _classify_identification_document(
        self,
        file_path: Path,
        schema_type: str,
        image_metadata: Dict[str, Any],
    ) -> Optional[Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]]:
        """PRIORITY 3.5: Detect identification documents in images (passport, ID, license).

        These should go to Person/ folder, not Media/.  Always runs OCR + KIE side
        effects (storing OCR confidence/language/text and any KIE result on self) so
        that downstream tiers can reuse them; returns the person tuple on a match or
        None to fall through.
        """
        if not (schema_type == "ImageObject" and self.ocr_available):
            return None
        # Extract text from image via OCR with confidence metadata.
        # Low-confidence results (e.g. blurry photos) must not trigger ID detection.
        _ocr_result = (
            extract_ocr_with_confidence(file_path, max_chars=0) if self.ocr_available else None
        )
        ocr_text = _ocr_result.text if _ocr_result else ""
        _ocr_conf = _ocr_result.confidence if _ocr_result else None
        _ocr_lang = _ocr_result.language if _ocr_result else None
        # Store for later persistence (consumed by _persist_to_graph_store).
        self._last_file_ocr_confidence = _ocr_conf
        self._last_file_detected_language = _ocr_lang
        # Cache OCR text so extract_text_from_image can reuse it.
        if ocr_text:
            self._last_file_ocr_text = ocr_text
        # Attempt KIE structured field extraction when OCR is reliable.
        if KIE_AVAILABLE and _ocr_conf is not None and _ocr_conf >= _OCR_CONFIDENCE_THRESHOLD:
            self._last_file_state["kie_result"] = extract_kie_fields(file_path)
        _id_conf_ok = _ocr_conf is None or _ocr_conf >= _OCR_CONFIDENCE_THRESHOLD
        if ocr_text and len(ocr_text) >= 30 and _id_conf_ok:
            ocr_lower = ocr_text.lower()
            # Check for identification document keywords
            id_keywords = [
                "passport",
                "driver license",
                "driver's license",
                "identification",
                "united states of america",
                "department of state",
                "nationality",
                "date of birth",
                "place of birth",
                "surname",
                "given names",
                "social security",
                "state id",
                "national id",
            ]
            if any(kw in ocr_lower for kw in id_keywords):
                print("  ✓ Identification document detected via OCR")
                people_names = []

                # Method 1: Parse passport MRZ (Machine Readable Zone)
                # Format: P<COUNTRY{SURNAME}<<{GIVEN_NAME}<...
                mrz_match = re.search(r"P<[A-Z]{3}([A-Z]+)<<([A-Z]+)<", ocr_text)
                if mrz_match:
                    surname = mrz_match.group(1).title()
                    given = mrz_match.group(2).title()
                    people_names = [f"{given} {surname}"]

                # Method 2: Look for name fields with values on next line or after colon
                # Passport format: "Surname\nLEDLIE" or "Surname/Nom\nLEDLIE"
                if not people_names:
                    # Find surname (all caps, standalone on line)
                    surname_match = re.search(
                        r"(?:surname|nom|apellidos)[/\w\s]*\n\s*([A-Z]{2,})\b",
                        ocr_text,
                        re.IGNORECASE,
                    )
                    given_match = re.search(
                        r"(?:given\s*names?|pr[ée]noms?|nombres)[/\w\s]*\n\s*([A-Z]{2,})\b",
                        ocr_text,
                        re.IGNORECASE,
                    )
                    if surname_match and given_match:
                        people_names = [
                            f"{given_match.group(1).title()} {surname_match.group(1).title()}"
                        ]

                # Method 3: General name extraction patterns
                if not people_names:
                    people_names = self.classifier.extract_people_names(ocr_text)

                if people_names:
                    print(f"  ✓ Person identified: {people_names[0]}")
                return (
                    "person",
                    "contacts",
                    schema_type,
                    ocr_text,
                    None,
                    people_names,
                    image_metadata,
                )

        return None

    def _classify_screenshot_ocr(
        self,
        file_path: Path,
        schema_type: str,
        image_metadata: Dict[str, Any],
    ) -> Optional[Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]]:
        """PRIORITY 4.5: Screenshot sub-classification via OCR + CLIP.

        Raw screenshots (e.g. "Screenshot 2025-*") that bypassed Priority 4 because
        classify_media_file returned None.  Try OCR first (reliable for screenshots
        with text), then CLIP for non-text images.  Returns a classification tuple
        for any screenshot-named image, or None for non-screenshots.
        """
        _stem_lower = file_path.stem.lower()
        if not (
            schema_type == "ImageObject"
            and (
                _stem_lower.startswith("screenshot")
                or "screen shot" in _stem_lower
                or (
                    "screenshot" in _stem_lower
                    and not re.match(
                        r"^(browser|terminal|code|docs|settings|product|chat|dashboard)_",
                        _stem_lower,
                    )
                )
            )
        ):
            return None

        screenshots_dict = self.category_paths["media"]["photos"]["screenshots"]

        # Step 1: OCR-based sub-classification (dashboard, terminal, etc.)
        ocr_result = _shared_classify_by_ocr(
            file_path,
            content_classifier=self.rename_analyzer.content_classifier,
        )
        if ocr_result:
            ocr_category, ocr_confidence, _ocr_scores, ocr_text = ocr_result
            if ocr_text:
                self._last_file_ocr_text = ocr_text
            if ocr_confidence < _SCREENSHOT_OCR_KEYWORD_THRESHOLD:
                print(
                    f"  ↪ Screenshot OCR low confidence ({ocr_confidence:.0%} < "
                    f"{_SCREENSHOT_OCR_KEYWORD_THRESHOLD:.0%}) — falling back to CLIP"
                )
            elif ocr_category in screenshots_dict:
                print(f"  ✓ Screenshot OCR sub-class: {ocr_category} ({ocr_confidence:.0%})")
                return (
                    "media",
                    f"photos_screenshots_{ocr_category}",
                    schema_type,
                    "",
                    None,
                    [],
                    image_metadata,
                )
            # OCR matched a non-screenshot Schema.org category — use it
            elif "_" in ocr_category:
                print(f"  ✓ Screenshot OCR reclassified: {ocr_category} ({ocr_confidence:.0%})")
                return (
                    ocr_category.split("_")[0],
                    ocr_category,
                    schema_type,
                    "",
                    None,
                    [],
                    image_metadata,
                )

        # Step 2: CLIP enhancement for images OCR couldn't classify
        enhanced = self.enhance_weak_image_classification(file_path, image_metadata)
        if enhanced:
            ecat, esubcat = enhanced
            # CLIP identified non-media content (e.g. game_assets) — use that
            if ecat not in ("media",):
                print(f"  ✓ Screenshot reclassified: {ecat}/{esubcat}")
                return (ecat, esubcat, schema_type, "", None, [], image_metadata)
            # CLIP identified a specific screenshot subcategory
            if "screenshots" in esubcat and esubcat != "photos_screenshots_other":
                print(f"  ✓ Screenshot CLIP sub-class: {esubcat}")
                return ("media", esubcat, schema_type, "", None, [], image_metadata)

        # Fallback: generic screenshot folder
        print("  ✓ Screenshot (unclassified)")
        return ("media", "photos_screenshots_other", schema_type, "", None, [], image_metadata)

    def _classify_photo_composition(
        self,
        file_path: Path,
        schema_type: str,
        image_metadata: Dict[str, Any],
    ) -> Optional[Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]]:
        """PRIORITY 5: Photo composition analysis (people / home interior).

        Returns a media/property_management tuple on a vision match, or None when
        vision is unavailable or no composition matched.
        """
        if not (schema_type == "ImageObject" and self.image_analyzer.vision_available):
            return None
        print("  Analyzing image content...")

        # Single CLIP pass yields both composition flags (people / home interior).
        has_people, is_property_mgmt, scores = self.image_analyzer.analyze_for_organization(
            file_path
        )

        if has_people:
            print("  ✓ Detected: Photo with people")
            if scores:
                top_categories = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:3]
                print(
                    f"  Top matches: {', '.join([f'{cat}: {score:.2%}' for cat, score in top_categories])}"  # noqa: E501
                )
            return ("media", "photos_social", schema_type, "", None, [], image_metadata)

        if is_property_mgmt:
            print("  ✓ Detected: Home interior without people")
            if scores:
                top_categories = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:3]
                print(
                    f"  Top matches: {', '.join([f'{cat}: {score:.2%}' for cat, score in top_categories])}"  # noqa: E501
                )
            return ("property_management", "other", schema_type, "", None, [], image_metadata)

        return None

    def _classify_by_content_and_kie(
        self,
        file_path: Path,
        schema_type: str,
        image_metadata: Dict[str, Any],
    ) -> Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]:
        """PRIORITY 6: Regular text extraction and content/KIE classification.

        Terminal tier — always returns a result tuple (the fall-through default for
        detect_file_category).
        """
        print("  Extracting content...")
        extracted_text = self.extract_text(file_path)

        if extracted_text:
            print(f"  Extracted {len(extracted_text)} characters")

            kie_result = self._last_file_state.get("kie_result")
            kie_classification = None
            if kie_result is not None:
                kie_classification = self.classifier.classify_with_kie(
                    kie_result,
                    extracted_text,
                    file_path.name,
                )
            if kie_classification is not None:
                category, subcategory, company_name, people_names = kie_classification
                print(f"  ✓ KIE classification: {category}/{subcategory}")
            else:
                category, subcategory, company_name, people_names = (
                    self.classifier.classify_content(extracted_text, file_path.name)
                )

            if company_name:
                print(f"  Detected company: {company_name}")
            if people_names:
                print(
                    f"  Detected people: {', '.join(people_names[:3])}{' ...' if len(people_names) > 3 else ''}"  # noqa: E501
                )
            print(f"  Classified as: {category}/{subcategory}")

            if schema_type == "ImageObject" and category != "uncategorized":
                category, subcategory = self._cross_check_with_clip(
                    file_path,
                    image_metadata,
                    category,
                    subcategory,
                    len(extracted_text),
                )
        else:
            print("  No text extracted, using filename")
            category, subcategory, company_name, people_names = self.classifier.classify_content(
                "", file_path.name
            )

        # Point C: last-resort enhancement for uncategorized images
        if category == "uncategorized" and schema_type == "ImageObject":
            enhanced = self.enhance_weak_image_classification(file_path, image_metadata)
            if enhanced:
                print(f"  ✓ Enhanced uncategorized: {enhanced[0]}/{enhanced[1]}")
                return (
                    enhanced[0],
                    enhanced[1],
                    schema_type,
                    extracted_text,
                    None,
                    [],
                    image_metadata,
                )

        return (
            category,
            subcategory,
            schema_type,
            extracted_text,
            company_name,
            people_names,
            image_metadata,
        )

    def generate_schema(self, file_path: Path, schema_type: str, extracted_text: str = "") -> Dict:
        """Generate Schema.org metadata for a file with extracted content."""
        stats = cached_stat(str(file_path))
        mime_type = self.enricher.detect_mime_type(str(file_path))
        file_url = f"https://localhost/files/{quote(file_path.name)}"
        actual_path = str(file_path.absolute())

        # Create generator based on type
        if schema_type == "ImageObject":
            generator = ImageGenerator(schema_type)
            generator.set_property("name", file_path.name, PropertyType.TEXT)
            generator.set_property("contentUrl", file_url, PropertyType.URL)
            generator.set_property("encodingFormat", mime_type or "image/png", PropertyType.TEXT)
            generator.set_property("description", f"{file_path.name}", PropertyType.TEXT)
        elif schema_type in ["DigitalDocument", "Article", SCHOLARLY_ARTICLE_SCHEMA_TYPE, "Report"]:
            generator = DocumentGenerator(schema_type)
            generator.set_property("name", file_path.name, PropertyType.TEXT)
            generator.set_property("description", f"{file_path.name}", PropertyType.TEXT)
            generator.set_property(
                "encodingFormat", mime_type or "application/octet-stream", PropertyType.TEXT
            )
            generator.set_property("url", file_url, PropertyType.URL)
            generator.set_property("contentSize", f"{stats.st_size}B", PropertyType.TEXT)
            research = self._last_file_state.get("research")
            if schema_type == SCHOLARLY_ARTICLE_SCHEMA_TYPE and research:
                _publisher_key, identifier, publisher_name, canonical_url = research
                try:
                    generator.set_property("identifier", identifier, PropertyType.TEXT)
                    generator.set_property("sameAs", canonical_url, PropertyType.URL)
                    generator.set_property(
                        "publisher",
                        {"@type": "Organization", "name": publisher_name},
                        PropertyType.OBJECT,
                    )
                except Exception as e:
                    print(f"  Warning: could not attach scholarly metadata: {e}")
        else:
            generator = DocumentGenerator()
            generator.set_property("name", file_path.name, PropertyType.TEXT)
            generator.set_property("description", f"{file_path.name}", PropertyType.TEXT)

        # Set dates
        try:
            generator.set_dates(
                created=datetime.fromtimestamp(stats.st_ctime),
                modified=datetime.fromtimestamp(stats.st_mtime),
            )
        except Exception:
            pass

        # Add extracted text as abstract/text property
        if extracted_text:
            try:
                # Truncate to reasonable length for schema
                text_preview = extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else "")
                generator.set_property("abstract", text_preview, PropertyType.TEXT)
                generator.set_property("text", extracted_text[:5000], PropertyType.TEXT)
            except Exception:
                pass

        # Add file path
        try:
            generator.set_property("filePath", actual_path, PropertyType.TEXT)
        except Exception:
            pass

        return generator.to_dict()

    def get_destination_path(
        self,
        file_path: Path,
        category: str,
        subcategory: str,
        company_name: Optional[str] = None,
        image_metadata: Optional[Dict] = None,
        people_names: Optional[List[str]] = None,
    ) -> Path:
        """
        Get the destination path for a file based on content category.

        Args:
            file_path: Path to the file
            category: Main category
            subcategory: Subcategory
            company_name: Optional company name for business/organization files
            image_metadata: Optional metadata for images (datetime, location)
            people_names: Optional list of people names for person-classified files

        Returns:
            Destination path for the file
        """
        # Special handling for filepath-based classification
        if category == "filepath":
            # subcategory contains the full path (e.g., "Technical/Python/MyProject")
            relative_path = subcategory
        # Special handling for media files with nested structure
        elif category == "media" and "_" in subcategory:
            # subcategory format: "photos_screenshots" or "photos_screenshots_browser"
            parts = subcategory.split("_", 1)  # Split into at most 2 parts
            if len(parts) == 2:
                media_type, media_subcat = parts
                if media_type in self.category_paths["media"]:
                    media_dict = self.category_paths["media"][media_type]
                    if isinstance(media_dict, dict):
                        # Check for 3-level nesting (e.g., screenshots_browser)
                        if "_" in media_subcat:
                            parent_key, child_key = media_subcat.split("_", 1)
                            parent_val = media_dict.get(parent_key)
                            if isinstance(parent_val, dict):
                                relative_path = parent_val.get(
                                    child_key,
                                    parent_val.get(
                                        "other",
                                        f"Media/{media_type.capitalize()}/{parent_key.capitalize()}",  # noqa: E501
                                    ),
                                )
                            else:
                                relative_path = media_dict.get(
                                    media_subcat,
                                    media_dict.get(
                                        "other", f"Media/{media_type.capitalize()}/Other"
                                    ),
                                )
                        else:
                            val = media_dict.get(media_subcat)
                            if isinstance(val, dict):
                                relative_path = val.get(
                                    "other",
                                    f"Media/{media_type.capitalize()}/{media_subcat.capitalize()}",
                                )
                            elif val:
                                relative_path = val
                            else:
                                relative_path = media_dict.get(
                                    "other", f"Media/{media_type.capitalize()}/Other"
                                )
                    else:
                        relative_path = media_dict
                else:
                    relative_path = "Media/Other"
            else:
                relative_path = "Media/Other"
        elif category in self.category_paths:
            if isinstance(self.category_paths[category], dict):
                if subcategory in self.category_paths[category]:
                    relative_path = self.category_paths[category][subcategory]
                else:
                    relative_path = self.category_paths[category].get(
                        "other", f"{category.capitalize()}/Other"
                    )
            else:
                relative_path = self.category_paths[category]
        else:
            relative_path = "Uncategorized"

        # Organization: Create entity-named subfolders under Organization/
        # Structure: Organization/{OrgName}/ for most types
        # Exception: Organization/Clients/{OrgName}/ for clients (nested subfolders)
        if category == "organization" and company_name:
            sanitized_company = self.classifier.sanitize_company_name(company_name)
            # Only create company subfolder if name is valid (not a sentence fragment)
            if sanitized_company:
                if subcategory == "clients":
                    # Clients get nested: Organization/Clients/{OrgName}/
                    relative_path = f"{relative_path}/{sanitized_company}"
                elif subcategory == "meeting_notes":
                    # Meeting notes get nested: Organization/{OrgName}/Meeting Notes/
                    relative_path = f"{relative_path}/{sanitized_company}/Meeting Notes"
                else:
                    # All other org types: Organization/{OrgName}/
                    relative_path = f"{relative_path}/{sanitized_company}"

        # Person: Create person-named subfolders under Person/
        # Structure: Person/{PersonName}/ for all types
        if category == "person" and people_names:
            # Use first person name as the folder name
            person_name = people_names[0] if people_names else "Unknown"
            sanitized_person = self.classifier.sanitize_company_name(person_name)
            # Only create person subfolder if name is valid
            if sanitized_person:
                relative_path = f"{relative_path}/{sanitized_person}"
            else:
                relative_path = f"{relative_path}/Unknown"
        elif category == "person" and not people_names:
            # Fallback for person category without extracted names
            relative_path = f"{relative_path}/Unknown"

        # Legacy: client files from business category with company name
        if category == "business" and subcategory == "clients" and company_name:
            sanitized_company = self.classifier.sanitize_company_name(company_name)
            # Only create company subfolder if name is valid
            if sanitized_company:
                relative_path = f"{relative_path}/{sanitized_company}"

        # Date-based organization for images (if enabled and metadata available)
        if self.organize_by_date and image_metadata and image_metadata.get("year"):
            year = image_metadata["year"]
            month = image_metadata["month"]
            relative_path = f"Photos/{year}/{month:02d}"

        # Location-based organization for images (if enabled and location available)
        elif self.organize_by_location and image_metadata and image_metadata.get("location_name"):
            # Clean location name for folder
            location = image_metadata["location_name"]
            # Take first part (usually city)
            city = location.split(",")[0].strip()
            # Sanitize for folder name
            safe_city = re.sub(r'[<>:"/\\|?*]', "", city)
            relative_path = f"Photos/Locations/{safe_city}"

        dest_dir = self.base_path / relative_path
        dest_dir.mkdir(parents=True, exist_ok=True)

        # Handle duplicate filenames
        dest_path = dest_dir / file_path.name
        if dest_path.exists() and dest_path != file_path:
            stem = file_path.stem
            suffix = file_path.suffix
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dest_path = dest_dir / f"{stem}_{timestamp}{suffix}"

        return dest_path

    def should_skip_file(self, file_path: Path) -> bool:
        """Check if file should be skipped."""
        skip_files = {".DS_Store", ".localized", "Thumbs.db", "desktop.ini"}
        skip_dirs = {"__pycache__", ".git", "node_modules", ".venv", "venv"}

        if file_path.name.startswith(".") and file_path.name not in {".gitignore", ".env.example"}:
            return True

        if file_path.name in skip_files:
            return True

        if any(skip_dir in file_path.parts for skip_dir in skip_dirs):
            return True

        return False

    def _persist_to_graph_store(
        self,
        file_path: Path,
        dest_path: Path,
        category: str,
        subcategory: str,
        schema: Dict,
        extracted_text: str,
        company_name: Optional[str],
        people_names: List[str],
        image_metadata: Optional[Dict],
        ocr_confidence: Optional[float] = None,
        detected_language: Optional[str] = None,
        kie_result=None,
    ) -> None:
        """
        Persist file and its relationships to the graph store with canonical IDs.

        This method creates:
        - File record with canonical_id (urn:sha256:{hash})
        - Category record with canonical_id (UUID v5 from name)
        - Company record with canonical_id (UUID v5 from name)
        - Person records with canonical_id (UUID v5 from name)
        - Location record with canonical_id (UUID v5 from name)
        - Relationships between file and entities
        """
        try:
            session = self.graph_store.get_session()

            # Get file stats
            stat = (
                cached_stat(str(file_path)) if file_path.exists() else cached_stat(str(dest_path))
            )

            # Merge KIE-extracted Schema.org properties into schema dict.
            kie_fields_json = None
            if kie_result is not None:
                try:
                    kie_schema = kie_result_to_schema_org(kie_result)
                    # Merge KIE properties without overwriting existing keys.
                    for k, v in kie_schema.items():
                        if k not in schema or k == "@type":
                            schema[k] = v
                    # Serialize raw fields for debugging/reprocessing.
                    kie_fields_json = {
                        cls: [{"value": f.value, "confidence": f.confidence} for f in fields]
                        for cls, fields in kie_result.fields.items()
                    }
                except Exception:
                    pass  # KIE merge failure must not block persistence

            # Add file to store (generates canonical_id automatically)
            file_record = self.graph_store.add_file(
                original_path=str(file_path),
                filename=file_path.name,
                session=session,
                current_path=str(dest_path),
                file_size=stat.st_size,
                mime_type=schema.get("encodingFormat"),
                schema_type=schema.get("@type"),
                schema_data=schema,
                extracted_text=extracted_text[:10000] if extracted_text else None,
                extracted_text_length=len(extracted_text) if extracted_text else 0,
                ocr_confidence=ocr_confidence,
                detected_language=detected_language,
                kie_fields=kie_fields_json,
                status=FileStatus.ORGANIZED,
                organized_at=datetime.now(),
            )

            file_id = file_record.id

            # Add category relationship
            self.graph_store.add_file_to_category(
                file_id=file_id,
                category_name=category,
                subcategory_name=subcategory,
                session=session,
            )

            # Add company relationship if detected
            if company_name:
                self.graph_store.add_file_to_company(
                    file_id=file_id,
                    company_name=company_name,
                    context="content_analysis",
                    session=session,
                )

            # Add people relationships if detected
            if people_names:
                for person_name in people_names:
                    self.graph_store.add_file_to_person(
                        file_id=file_id, person_name=person_name, role="mentioned", session=session
                    )

            # Add location if available from image metadata
            if image_metadata and image_metadata.get("location"):
                location_info = image_metadata["location"]
                self.graph_store.add_file_to_location(
                    file_id=file_id,
                    location_name=location_info.get("display_name", "Unknown"),
                    latitude=location_info.get("latitude"),
                    longitude=location_info.get("longitude"),
                    city=location_info.get("city"),
                    state=location_info.get("state"),
                    country=location_info.get("country"),
                    location_type="captured_at",
                    session=session,
                )

            session.commit()
            session.close()

        except Exception as e:
            print(f"  ⚠ Graph store error (non-fatal): {e}")

    def _maybe_rename_image(self, file_path: Path, dry_run: bool) -> Path:
        """Rename generic image files using content analysis before sorting.

        When *not* dry-run, physically renames the file and returns the
        new path.  In dry-run mode the file stays on disk but the
        proposed new path is returned so that filename-pattern
        classification sees the descriptive name.  Callers that need to
        read file contents should use the original path stored in
        ``result['source']``.
        """
        if not is_generic_filename(file_path.name):
            return file_path

        if file_path.suffix.lower() not in IMAGE_EXTENSIONS_WIDE:
            return file_path

        result = self.rename_analyzer.analyze_image(file_path)

        new_name = result.get("new_name")
        if not new_name or result.get("status") != ProcessingStatus.PENDING:
            return file_path

        conf = result.get("confidence")
        conf_str = f" ({conf:.0%})" if conf is not None else ""
        new_path = resolve_collision(file_path.parent / new_name)

        if dry_run:
            print(f"  → Would rename: {file_path.name} → {new_path.name}{conf_str}")
            return new_path

        file_path.rename(new_path)
        print(f"  ✓ Renamed: {file_path.name} → {new_path.name}{conf_str}")
        return new_path

    def organize_file(self, file_path: Path, dry_run: bool = False, force: bool = False) -> Dict:
        """
        Organize a single file based on content.

        Args:
            file_path: Path to the file
            dry_run: If True, don't actually move files
            force: If True, re-organize even if already in correct location

        Returns:
            Dictionary with organization details
        """
        result = {
            "source": str(file_path),
            "status": "skipped",
            "reason": None,
            "destination": None,
            "schema": None,
            "extracted_text_length": 0,
        }

        if self.should_skip_file(file_path):
            result["reason"] = "system_file"
            self.stats["skipped"] += 1
            return result

        if not file_path.is_file():
            result["reason"] = "not_file"
            self.stats["skipped"] += 1
            return result

        try:
            # Rename generic image files (screenshots, IMG_, etc.) before classification.
            # In dry-run the file stays on disk at file_path but renamed_path
            # carries the descriptive name for pattern matching.
            renamed_path = self._maybe_rename_image(file_path, dry_run)
            display_path = renamed_path if renamed_path != file_path else None
            physical_path = renamed_path if not dry_run else file_path

            # Detect category: physical_path for content reading,
            # display_path (renamed name) for filename-pattern matching.
            (
                category,
                subcategory,
                schema_type,
                extracted_text,
                company_name,
                people_names,
                image_metadata,
            ) = self.detect_file_category(physical_path, display_path=display_path)
            result["extracted_text_length"] = len(extracted_text)
            result["company_name"] = company_name
            result["people_names"] = people_names
            result["image_metadata"] = image_metadata

            # Handle skip category (duplicates, etc.)
            if category == "skip":
                result["status"] = "skipped"
                result["reason"] = subcategory  # e.g., 'duplicate'
                self.stats["skipped"] += 1
                return result

            # Generate schema with extracted content.
            # Use physical_path (current path on disk) since the file may have
            # been renamed by _maybe_rename_image before reaching this point.
            schema = self.generate_schema(physical_path, schema_type, extracted_text)

            # Validate schema
            validation_report = self.validator.validate(schema)

            # Get destination path (with optional date/location organization for images)
            # Use renamed_path so the destination carries the descriptive filename.
            dest_path = self.get_destination_path(
                renamed_path, category, subcategory, company_name, image_metadata, people_names
            )

            # Skip if already in the right place (unless force=True)
            if physical_path == dest_path and not force:
                result["status"] = "already_organized"
                result["destination"] = str(dest_path)
                result["schema"] = schema
                result["category"] = category
                result["subcategory"] = subcategory
                self.stats["already_organized"] += 1
                return result

            # Move file if not dry run
            if not dry_run:
                shutil.move(str(physical_path), str(dest_path))

                # Register schema
                schema["url"] = f"file://{dest_path.absolute()}"
                metadata = {
                    "category": category,
                    "subcategory": subcategory,
                    "organized_date": datetime.now().isoformat(),
                    "is_valid": validation_report.is_valid(),
                    "has_extracted_text": bool(extracted_text),
                }
                if company_name:
                    metadata["company_name"] = company_name

                self.registry.register(str(dest_path), schema, metadata=metadata)

                # Persist to database with canonical IDs
                if self.graph_store:
                    self._persist_to_graph_store(
                        file_path=file_path,
                        dest_path=dest_path,
                        category=category,
                        subcategory=subcategory,
                        schema=schema,
                        extracted_text=extracted_text,
                        company_name=company_name,
                        people_names=people_names,
                        image_metadata=image_metadata,
                        ocr_confidence=self._last_file_ocr_confidence,
                        detected_language=self._last_file_detected_language,
                        kie_result=self._last_file_state.get("kie_result"),
                    )

            result["status"] = "organized" if not dry_run else "would_organize"
            result["destination"] = str(dest_path)
            result["schema"] = schema
            result["category"] = category
            result["subcategory"] = subcategory
            result["is_valid"] = validation_report.is_valid()

            self.stats["organized"] += 1
            self.stats[f"{category}_{subcategory}"] += 1

        except Exception as e:
            result["status"] = "error"
            result["reason"] = str(e)
            self.stats["errors"] += 1
            print(f"  ✗ Error: {e}")

        return result

    def scan_directory(self, directory: Path) -> List[Path]:
        """Scan directory for files to organize."""
        files = []
        try:
            for item in directory.rglob("*"):
                if item.is_file() and not self.should_skip_file(item):
                    files.append(item)
        except PermissionError:
            print(f"Permission denied: {directory}")
        return files

    def organize_directories(
        self, source_dirs: List[str], dry_run: bool = False, limit: int = None, force: bool = False
    ) -> Dict:
        """
        Organize files from multiple source directories.

        Args:
            source_dirs: List of directory paths to organize
            dry_run: If True, simulate organization without moving files
            limit: Maximum number of files to process (for testing)
            force: If True, re-organize files even if already in correct location

        Returns:
            Dictionary with organization results
        """
        results = []

        print(f"\n{'='*60}")
        print(f"Content-Based File Organization {'(DRY RUN)' if dry_run else ''}")
        print(f"{'='*60}\n")

        if not self.ocr_available:
            print("⚠️  WARNING: OCR libraries not available")
            print("   Install with: pip install python-doctr[torch] Pillow pypdf")
            print("   Content classification will be limited to filenames\n")

        # Scan all directories
        all_files = []
        for source_dir in source_dirs:
            source_path = Path(source_dir).expanduser()
            if source_path.exists():
                print(f"Scanning: {source_path}")
                files = self.scan_directory(source_path)
                all_files.extend(files)
                print(f"  Found {len(files)} files")
            else:
                print(f"Directory not found: {source_path}")

        if limit:
            all_files = all_files[:limit]
            print(f"\n⚠️  Processing limited to first {limit} files for testing\n")

        print(f"\nTotal files to process: {len(all_files)}\n")

        # Organize each file
        for i, file_path in enumerate(all_files, 1):
            print(f"[{i}/{len(all_files)}] Processing: {file_path.name}")
            result = self.organize_file(file_path, dry_run=dry_run, force=force)
            results.append(result)

            if result["status"] == "organized" or result["status"] == "would_organize":
                print(f"  → {result['destination']}")
            elif result["status"] == "error":
                print(f"  ✗ Error: {result['reason']}")

        # Generate summary
        summary = {
            "total_files": len(all_files),
            "organized": self.stats["organized"],
            "already_organized": self.stats["already_organized"],
            "skipped": self.stats["skipped"],
            "errors": self.stats["errors"],
            "dry_run": dry_run,
            "results": results,
            "registry_stats": self.registry.get_statistics() if not dry_run else None,
        }

        return summary

    def print_summary(self, summary: Dict):
        """Print organization summary."""
        print(f"\n{'='*60}")
        print("Organization Summary")
        print(f"{'='*60}\n")

        print(f"Total files processed: {summary['total_files']}")
        print(f"Successfully organized: {summary['organized']}")
        print(f"Already organized: {summary['already_organized']}")
        print(f"Skipped: {summary['skipped']}")
        print(f"Errors: {summary['errors']}")

        if summary["dry_run"]:
            print("\n⚠️  This was a DRY RUN - no files were moved")

        # Category breakdown
        print(f"\n{'='*60}")
        print("Category Breakdown")
        print(f"{'='*60}\n")

        category_stats = defaultdict(int)
        for result in summary["results"]:
            if result.get("category"):
                category_stats[result["category"]] += 1

        for category, count in sorted(category_stats.items()):
            print(f"{category.capitalize()}: {count} files")

        # OCR stats
        ocr_count = sum(1 for r in summary["results"] if r.get("extracted_text_length", 0) > 0)
        print(f"\n{'='*60}")
        print("Content Extraction Stats")
        print(f"{'='*60}\n")
        print(f"Files with extracted text: {ocr_count}/{summary['total_files']}")

        # Company detection stats
        company_files = [r for r in summary["results"] if r.get("company_name")]
        if company_files:
            print(f"\n{'='*60}")
            print("Detected Companies")
            print(f"{'='*60}\n")
            company_counts = defaultdict(int)
            for result in company_files:
                company_counts[result["company_name"]] += 1

            print(f"Total files with detected companies: {len(company_files)}")
            print("\nCompanies found:")
            for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
                print(f"  {company}: {count} files")

        if summary.get("registry_stats"):
            print(f"\n{'='*60}")
            print("Schema Registry")
            print(f"{'='*60}\n")
            stats = summary["registry_stats"]
            print(f"Total schemas: {stats['total_schemas']}")
            print(f"Types: {', '.join(stats['types'])}")

        # Cost tracking summary
        if self.cost_calculator:
            self._print_cost_summary()

    def _print_cost_summary(self):
        """Print cost and ROI summary from the cost calculator."""
        if not self.cost_calculator:
            return

        print(f"\n{'='*60}")
        print("Cost & ROI Analysis")
        print(f"{'='*60}\n")

        cost_summary = self.cost_calculator.calculate_total_cost()
        roi_summary = self.cost_calculator.calculate_total_roi()

        print(f"Total Processing Cost:     ${cost_summary['total_cost']:.4f}")
        print(f"Total Files Processed:     {cost_summary['total_files_processed']:,}")
        print(f"Avg Cost per File:         ${cost_summary['avg_cost_per_file']:.6f}")
        print(f"Total Processing Time:     {cost_summary['total_processing_time_sec']:.1f}s")

        print(f"\nEstimated Value Generated: ${roi_summary['total_value']:.2f}")
        roi_pct = roi_summary["overall_roi_percentage"]
        roi_str = f"{roi_pct:.0f}%" if roi_pct != float("inf") else "∞"
        print(f"Overall ROI:               {roi_str}")
        print(f"Manual Hours Saved:        {roi_summary['total_manual_hours_saved']:.1f} hours")

        # Per-feature breakdown (top 5 by usage)
        feature_costs = cost_summary.get("feature_breakdown", {})
        if feature_costs:
            print(f"\n{'Feature':<25} {'Cost':>10} {'Files':>10}")
            print("-" * 50)
            sorted_features = sorted(
                feature_costs.items(), key=lambda x: x[1]["total_files_processed"], reverse=True
            )
            for feature_name, data in sorted_features[:7]:
                if data["total_invocations"] > 0:
                    print(
                        f"{feature_name:<25} ${data['total_cost']:>9.4f} {data['total_files_processed']:>10,}"  # noqa: E501
                    )

        # Show recommendations if any critical issues
        recommendations = self.cost_calculator.get_optimization_recommendations()
        critical_recs = [r for r in recommendations if r["severity"] in ("critical", "high")]
        if critical_recs:
            print("\n⚠️  Optimization Recommendations:")
            for rec in critical_recs[:3]:
                print(f"   • {rec['message']}")

    def get_cost_report(self) -> Optional[Dict[str, Any]]:
        """
        Get the full cost and ROI report.

        Returns:
            Cost report dictionary or None if cost tracking is disabled
        """
        if not self.cost_calculator:
            return None
        return self.cost_calculator.generate_report()

    def save_cost_report(self, output_path: str = None):
        """
        Save the cost report to a JSON file.

        Args:
            output_path: Path to save the report (auto-generated if None)
        """
        if not self.cost_calculator:
            print("Cost tracking is not enabled")
            return

        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"results/cost_report_{timestamp}.json"

        self.cost_calculator.generate_report(output_path)
        print(f"Cost report saved to: {output_path}")

    def save_report(self, summary: Dict, output_path: str = None):
        """Save detailed organization report to JSON."""
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"results/content_organization_report_{timestamp}.json"

        with open(output_path, "w") as f:
            json.dump(summary, f, indent=2, default=str)

        print(f"\nDetailed report saved to: {output_path}")


def main():
    """Main entry point."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Organize files by content using OCR and Schema.org metadata"
    )
    parser.add_argument(
        "--dry-run", action="store_true", help="Simulate organization without moving files"
    )
    parser.add_argument(
        "--base-path",
        default="~/Documents",
        help="Base path for organized files (default: ~/Documents)",
    )
    parser.add_argument(
        "--sources",
        nargs="+",
        default=["~/Desktop", "~/Downloads"],
        help="Source directories to organize (default: ~/Desktop ~/Downloads)",
    )
    parser.add_argument("--report", help="Path to save detailed JSON report")
    parser.add_argument("--limit", type=int, help="Limit number of files to process (for testing)")
    parser.add_argument(
        "--no-cost-tracking", action="store_true", help="Disable cost and ROI tracking"
    )
    parser.add_argument(
        "--cost-report",
        nargs="?",
        const="results/cost_report.json",
        default="results/cost_report.json",
        help=(
            "Path to save cost/ROI report (default: results/cost_report.json, "
            "use --no-cost-tracking to disable)"
        ),
    )
    parser.add_argument(
        "--check-deps",
        action="store_true",
        help="Run system health check and show feature availability",
    )
    parser.add_argument(
        "--skip-health-check", action="store_true", help="Skip startup health check"
    )
    parser.add_argument(
        "--sentry-dsn", help="Sentry DSN for error tracking (or set SENTRY_DSN env var)"
    )
    parser.add_argument("--no-sentry", action="store_true", help="Disable Sentry error tracking")
    parser.add_argument(
        "--db-path",
        default="results/file_organization.db",
        help=(
            "Path to SQLite database for persistent storage "
            "(default: results/file_organization.db)"
        ),
    )
    parser.add_argument(
        "--no-db",
        action="store_true",
        help="Disable database persistence (use in-memory registry only)",
    )
    parser.add_argument(
        "--run-migration",
        action="store_true",
        help="Run database migration to add canonical_id columns to existing records",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Force re-organization of all files, even if already in correct location",
    )

    args = parser.parse_args()

    # Initialize Sentry error tracking (before any other operations)
    if not args.no_sentry and ERROR_TRACKING_AVAILABLE:
        # Priority: CLI arg > FILE_SYSTEM_SENTRY_DSN > SENTRY_DSN
        sentry_dsn = (
            args.sentry_dsn
            or os.environ.get("FILE_SYSTEM_SENTRY_DSN")
            or os.environ.get("SENTRY_DSN")
        )
        if sentry_dsn:
            os.environ["SENTRY_DSN"] = sentry_dsn
        sentry_enabled = init_sentry()
        if sentry_enabled:
            print("✓ Sentry error tracking enabled")
    else:
        sentry_enabled = False

    # Run system health check
    if args.check_deps:
        from health_check import check_system

        check_system(verbose=True)
        return

    if not args.skip_health_check:
        from health_check import SystemHealthChecker

        checker = SystemHealthChecker().run_all_checks()
        checker.print_status()

    # Run migration if requested
    if args.run_migration:
        if GRAPH_STORE_AVAILABLE:
            from storage.migration import run_migration

            print(f"\n{'='*60}")
            print("Running ID Generation Migration")
            print(f"{'='*60}\n")
            run_migration(args.db_path)
            print("\nMigration complete. Canonical IDs have been generated for existing records.")
            return
        else:
            print("Error: GraphStore not available. Cannot run migration.")
            return

    # Create organizer with database path
    db_path = None if args.no_db else args.db_path
    organizer = ContentBasedFileOrganizer(
        base_path=args.base_path, enable_cost_tracking=not args.no_cost_tracking, db_path=db_path
    )

    # Organize directories
    summary = organizer.organize_directories(
        source_dirs=args.sources, dry_run=args.dry_run, limit=args.limit, force=args.force
    )

    # Print summary
    organizer.print_summary(summary)

    # Save reports
    if args.report or not args.dry_run:
        organizer.save_report(summary, args.report)

    # Save cost report if tracking was enabled
    if not args.no_cost_tracking and organizer.cost_calculator:
        organizer.save_cost_report(args.cost_report)

    # Update _site directory with latest HTML files
    if not args.dry_run:
        import subprocess
        from pathlib import Path

        script_path = Path(__file__).parent / "copy_to_site.sh"
        if script_path.exists():
            try:
                subprocess.run([str(script_path)], check=True, capture_output=True)
                print("\n✓ Updated _site directory with latest HTML files")
            except subprocess.CalledProcessError:
                print("\n⚠ Failed to update _site directory")


if __name__ == "__main__":
    main()
