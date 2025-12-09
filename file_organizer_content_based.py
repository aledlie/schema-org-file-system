#!/usr/bin/env python3
"""
Content-Based Intelligent File Organizer using Schema.org metadata and OCR.

Organizes files based on their actual content rather than just file type.
Uses OCR to extract text from images and PDFs, then classifies by content.
"""

import sys
import os
import shutil
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from collections import defaultdict
from urllib.parse import quote

# OCR and PDF imports
try:
    import pytesseract
    from PIL import Image
    import pypdf
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True

    # HEIC support for OCR
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
    except ImportError:
        pass  # HEIC support optional for OCR
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not available. Install pytesseract, Pillow, pypdf, pdf2image")

# Word document imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install python-docx")

# Excel imports
try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Warning: openpyxl not available. Install openpyxl")

# Add src directory to path
sys.path.insert(0, '/Users/alyshialedlie/schema-org-file-system/src')

from generators import (
    DocumentGenerator,
    ImageGenerator,
    VideoGenerator,
    AudioGenerator,
    CodeGenerator,
    DatasetGenerator,
    ArchiveGenerator
)
from base import PropertyType
from enrichment import MetadataEnricher
from validator import SchemaValidator
from integration import SchemaRegistry

# Image content analysis imports
try:
    from transformers import CLIPProcessor, CLIPModel
    import torch
    import cv2
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False
    print("Warning: Vision libraries not available. Install transformers, torch, opencv-python")

# Image metadata imports
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    import piexif
    from geopy.geocoders import Nominatim
    from geopy.exc import GeocoderTimedOut, GeocoderServiceError
    METADATA_AVAILABLE = True

    # HEIC support
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
        HEIC_AVAILABLE = True
    except ImportError:
        HEIC_AVAILABLE = False
except ImportError:
    METADATA_AVAILABLE = False
    print("Warning: Metadata libraries not available. Install piexif, geopy")


class ContentClassifier:
    """Classifies document content into categories."""

    def __init__(self):
        """Initialize classifier with keyword patterns."""
        # Company name patterns
        self.company_patterns = [
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+LLC\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+L\.L\.C\.\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Inc\.?\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Incorporated\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Corp\.?\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Corporation\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Company\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Co\.\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Ltd\.?\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Limited\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+LLP\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+L\.L\.P\.\b',
        ]

        # People name patterns - look for common name patterns
        self.people_patterns = [
            # Name with document type indicators
            r'\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\s+(?:Resume|CV|Cover Letter)\b',
            r'\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\s+(?:Portfolio|Biography|Bio)\b',
            # Field labels followed by names
            r'\b(?:Name|Contact|From|To|Attn|Author|Client|Patient|Student):\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b',
            # Email signatures (name before email)
            r'\b([A-Z][a-z]+\s+[A-Z][a-z]+)\s+<[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}>',
            # Name in "Prepared by/for" statements
            r'\b(?:Prepared|Written|Submitted|Signed)\s+(?:by|for):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b',
            # Name followed by credentials (MD, PhD, Esq, etc.)
            r'\b([A-Z][a-z]+\s+[A-Z][a-z]+),?\s+(?:MD|PhD|Esq|DDS|CPA|MBA|JD|RN)\b',
            # Mr./Mrs./Ms./Dr. followed by name
            r'\b(?:Mr|Mrs|Ms|Dr|Prof)\.?\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b',
            # Name in meeting notes format
            r'\b(?:Attendee|Participant|Speaker|Presenter):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b',
        ]

        self.patterns = {
            'legal': {
                'keywords': [
                    'contract', 'agreement', 'terms', 'conditions', 'legal', 'attorney',
                    'law', 'litigation', 'plaintiff', 'defendant', 'court', 'settlement',
                    'lease', 'deed', 'will', 'testament', 'power of attorney', 'notary',
                    'amendment', 'exhibit', 'whereas', 'party', 'parties', 'executed',
                    'operating agreement', 'llc', 'corporation', 'bylaws', 'articles'
                ],
                'subcategories': {
                    'contracts': ['contract', 'agreement', 'terms', 'subscription', 'saas'],
                    'real_estate': ['lease', 'deed', 'property', 'real estate', 'mortgage'],
                    'corporate': ['llc', 'corporation', 'operating agreement', 'bylaws', 'articles', 'formation'],
                    'other': []
                }
            },
            'financial': {
                'keywords': [
                    'invoice', 'receipt', 'tax', 'irs', 'payment', 'bill', 'billing',
                    'statement', 'account', 'balance', 'transaction', 'credit', 'debit',
                    'bank', 'finance', 'loan', 'interest', '1098', '1099', 'w-2', 'w2',
                    'federal', 'state return', 'refund', 'revenue', 'expense', 'budget',
                    'investment', 'portfolio', 'ein', 'employer identification'
                ],
                'subcategories': {
                    'tax': ['tax', 'irs', '1098', '1099', 'w-2', 'w2', 'federal', 'state return'],
                    'invoices': ['invoice', 'bill', 'billing', 'payment'],
                    'statements': ['statement', 'account', 'balance', 'transaction'],
                    'other': []
                }
            },
            'business': {
                'keywords': [
                    'proposal', 'pitch', 'business plan', 'strategy', 'marketing',
                    'presentation', 'deck', 'startup', 'company', 'venture', 'investor',
                    'growth', 'revenue model', 'unit economics', 'expansion', 'rfp',
                    'guidelines', 'program', 'service package', 'pricing', 'client',
                    'customer', 'vendor', 'supplier', 'partner'
                ],
                'subcategories': {
                    'planning': ['business plan', 'strategy', 'expansion', 'growth'],
                    'marketing': ['marketing', 'pricing', 'service package', 'pitch', 'deck'],
                    'proposals': ['proposal', 'rfp', 'guidelines'],
                    'clients': ['client', 'customer', 'llc', 'inc', 'corp', 'company'],
                    'other': []
                }
            },
            'personal': {
                'keywords': [
                    'resume', 'cv', 'cover letter', 'curriculum vitae', 'employment',
                    'personal', 'identification', 'passport', 'driver license', 'ssn',
                    'birth certificate', 'marriage', 'divorce', 'diploma', 'transcript',
                    'reference', 'recommendation'
                ],
                'subcategories': {
                    'employment': ['resume', 'cv', 'cover letter', 'employment', 'reference'],
                    'identification': ['passport', 'driver license', 'ssn', 'id'],
                    'certificates': ['birth certificate', 'marriage', 'divorce', 'diploma'],
                    'other': []
                }
            },
            'medical': {
                'keywords': [
                    'medical', 'health', 'doctor', 'patient', 'prescription', 'diagnosis',
                    'treatment', 'hospital', 'clinic', 'insurance claim', 'hipaa',
                    'vaccination', 'immunization', 'lab results', 'pharmacy'
                ],
                'subcategories': {
                    'records': ['medical record', 'patient', 'diagnosis', 'treatment'],
                    'insurance': ['insurance', 'claim', 'coverage'],
                    'prescriptions': ['prescription', 'pharmacy', 'medication'],
                    'other': []
                }
            },
            'property': {
                'keywords': [
                    'property management', 'tenant', 'landlord', 'rent', 'rental',
                    'maintenance', 'repair', 'inspection', 'utilities', 'hoa'
                ],
                'subcategories': {
                    'leases': ['lease', 'tenant', 'landlord', 'rent', 'rental'],
                    'maintenance': ['maintenance', 'repair', 'inspection'],
                    'other': []
                }
            },
            'education': {
                'keywords': [
                    'course', 'syllabus', 'lecture', 'assignment', 'homework', 'exam',
                    'grade', 'transcript', 'diploma', 'degree', 'certificate', 'university',
                    'college', 'school', 'research paper', 'thesis', 'dissertation'
                ],
                'subcategories': {
                    'coursework': ['course', 'syllabus', 'lecture', 'assignment'],
                    'research': ['research', 'paper', 'thesis', 'dissertation'],
                    'records': ['transcript', 'diploma', 'degree', 'certificate'],
                    'other': []
                }
            },
            'technical': {
                'keywords': [
                    'code', 'software', 'development', 'programming', 'api', 'database',
                    'documentation', 'technical', 'specification', 'architecture', 'design',
                    'system', 'infrastructure', 'deployment', 'configuration'
                ],
                'subcategories': {
                    'documentation': ['documentation', 'spec', 'specification', 'readme'],
                    'architecture': ['architecture', 'design', 'system', 'infrastructure'],
                    'other': []
                }
            },
            'creative': {
                'keywords': [
                    'design', 'graphic', 'illustration', 'artwork', 'photo', 'image',
                    'screenshot', 'mockup', 'prototype', 'wireframe', 'brand', 'logo'
                ],
                'subcategories': {
                    'design': ['design', 'mockup', 'wireframe', 'prototype'],
                    'branding': ['brand', 'logo', 'identity'],
                    'photos': ['photo', 'photography', 'image'],
                    'other': []
                }
            }
        }

    def extract_company_names(self, text: str) -> List[str]:
        """
        Extract company names from text using regex patterns.

        Returns:
            List of detected company names
        """
        companies = []
        for pattern in self.company_patterns:
            matches = re.findall(pattern, text)
            companies.extend(matches)

        # Remove duplicates and clean up
        unique_companies = []
        seen = set()
        for company in companies:
            # Clean up whitespace
            clean = ' '.join(company.split())
            # Skip if too short or already seen
            if len(clean) > 2 and clean.lower() not in seen:
                seen.add(clean.lower())
                unique_companies.append(clean)

        return unique_companies

    def extract_people_names(self, text: str) -> List[str]:
        """
        Extract people names from text using regex patterns.

        Returns:
            List of detected people names
        """
        people = []
        for pattern in self.people_patterns:
            matches = re.findall(pattern, text)
            # Pattern can return tuples (first, last) or single strings
            for match in matches:
                if isinstance(match, tuple):
                    # Join tuple elements (e.g., first name + last name)
                    full_name = ' '.join([m for m in match if m])
                else:
                    full_name = match
                people.append(full_name)

        # Remove duplicates and clean up
        unique_people = []
        seen = set()
        for person in people:
            # Clean up whitespace
            clean = ' '.join(person.split())
            # Skip if too short or already seen
            if len(clean) > 2 and clean.lower() not in seen:
                seen.add(clean.lower())
                unique_people.append(clean)

        return unique_people

    def extract_person_company_relationships(self, text: str) -> Dict[str, str]:
        """
        Extract relationships between people and companies from text.
        Uses Schema.org-style connections (Person worksFor/memberOf Organization).

        Returns:
            Dictionary mapping person names to company names
        """
        relationships = {}

        # Patterns for person-company relationships
        relationship_patterns = [
            # "John Doe at Company LLC"
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+(?:at|from)\s+([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))',
            # "John Doe, CEO of Company LLC"
            r'([A-Z][a-z]+\s+[A-Z][a-z]+),?\s+(?:CEO|CFO|CTO|COO|President|Director|Manager|Founder)\s+(?:of|at)\s+([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))',
            # "Company LLC - Contact: John Doe"
            r'([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))\s*[-:]\s*(?:Contact|Representative):\s*([A-Z][a-z]+\s+[A-Z][a-z]+)',
            # "John Doe (Company LLC)"
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+\(([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))\)',
            # Email pattern: john.doe@company.com -> John Doe at Company
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+<[a-zA-Z0-9._%+-]+@([a-zA-Z0-9.-]+)\.[a-zA-Z]{2,}>',
        ]

        for pattern in relationship_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if len(match) == 2:
                    person, company = match
                    # Clean up
                    person_clean = ' '.join(person.split())
                    company_clean = ' '.join(company.split())

                    # For email domains, capitalize company name
                    if '@' in text and '.' in company_clean and len(company_clean.split('.')) >= 2:
                        # This is likely a domain name, extract company name
                        domain_parts = company_clean.split('.')
                        if domain_parts[0].lower() not in ['gmail', 'yahoo', 'hotmail', 'outlook', 'mail']:
                            company_clean = domain_parts[0].capitalize()

                    # Store relationship (person -> company)
                    if len(person_clean) > 2 and len(company_clean) > 2:
                        relationships[person_clean] = company_clean

        return relationships

    def sanitize_company_name(self, company_name: str) -> str:
        """
        Sanitize company name for use in folder names.

        Returns:
            Sanitized folder name
        """
        # Remove special characters that aren't allowed in folder names
        sanitized = re.sub(r'[<>:"/\\|?*]', '', company_name)
        # Replace multiple spaces with single space
        sanitized = ' '.join(sanitized.split())
        # Limit length
        if len(sanitized) > 50:
            sanitized = sanitized[:50].strip()
        return sanitized

    def classify_content(self, text: str, filename: str = "") -> Tuple[str, str, Optional[str], List[str]]:
        """
        Classify content based on extracted text.
        Uses Schema.org person-company relationships to improve categorization.

        Returns:
            Tuple of (category, subcategory, company_name, people_names)
        """
        if not text:
            return ('uncategorized', 'other', None, [])

        text_lower = text.lower()
        filename_lower = filename.lower()
        combined = f"{text_lower} {filename_lower}"

        # Extract company names and people names
        company_names = self.extract_company_names(text)
        primary_company = company_names[0] if company_names else None

        people_names = self.extract_people_names(text)

        # Extract person-company relationships (Schema.org connections)
        person_company_relationships = self.extract_person_company_relationships(text)

        # Prioritize company from person-company relationships over direct extraction
        # Relationships tend to be more accurate as they include context
        if person_company_relationships:
            # Get the first relationship's company
            relationship_company = next(iter(person_company_relationships.values()))

            # Check if relationship company has proper legal suffix
            has_legal_suffix = any(relationship_company.endswith(suffix)
                                  for suffix in ['LLC', 'Inc.', 'Inc', 'Corp.', 'Corp',
                                                'Ltd.', 'Ltd', 'LLP', 'L.L.C.', 'L.L.P.'])

            # Prefer relationship company if it has legal suffix or we don't have a primary company
            if has_legal_suffix or not primary_company:
                primary_company = relationship_company
            # Or if the relationship company is much cleaner (shorter and no weird prefixes)
            elif primary_company and 'CEO' not in relationship_company and 'at' not in relationship_company:
                if len(relationship_company) < len(primary_company) * 0.8:
                    primary_company = relationship_company

        # Fallback: If we found people but no company, check relationships again
        if people_names and not primary_company and person_company_relationships:
            # Try to find a company for the first person mentioned
            for person in people_names:
                if person in person_company_relationships:
                    primary_company = person_company_relationships[person]
                    break

        # Score each category
        scores = defaultdict(int)
        category_subcats = {}

        for category, data in self.patterns.items():
            for keyword in data['keywords']:
                count = combined.count(keyword.lower())
                if count > 0:
                    scores[category] += count

                    # Track which subcategory keywords matched
                    for subcat, subcat_keywords in data['subcategories'].items():
                        if any(sk.lower() in combined for sk in subcat_keywords):
                            if category not in category_subcats:
                                category_subcats[category] = defaultdict(int)
                            category_subcats[category][subcat] += count

        if not scores:
            return ('uncategorized', 'other', primary_company, people_names)

        # Get category with highest score
        best_category = max(scores.items(), key=lambda x: x[1])[0]

        # Get subcategory with highest score for this category
        if best_category in category_subcats:
            subcat_scores = category_subcats[best_category]
            if subcat_scores:
                best_subcategory = max(subcat_scores.items(), key=lambda x: x[1])[0]
            else:
                best_subcategory = 'other'
        else:
            best_subcategory = 'other'

        # If we detected a company (either directly or via person relationship)
        # and it's business-related, use clients subcategory
        if primary_company and best_category == 'business':
            best_subcategory = 'clients'

        return (best_category, best_subcategory, primary_company, people_names)


class ImageMetadataParser:
    """Parses image metadata including EXIF, GPS, and timestamps."""

    def __init__(self):
        """Initialize the metadata parser."""
        self.metadata_available = METADATA_AVAILABLE
        self.geocoder = None

        if self.metadata_available:
            try:
                # Initialize geocoder with a user agent
                self.geocoder = Nominatim(user_agent="file_organizer_v1.0", timeout=5)
            except Exception as e:
                print(f"Warning: Could not initialize geocoder: {e}")
                self.geocoder = None

    def extract_exif_data(self, image_path: Path) -> Dict[str, Any]:
        """
        Extract EXIF data from an image.

        Returns:
            Dictionary with EXIF data
        """
        if not self.metadata_available:
            return {}

        try:
            image = Image.open(image_path)
            exif_data = {}

            # Get EXIF data
            exif = image._getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    exif_data[tag] = value

            return exif_data

        except Exception as e:
            print(f"  EXIF extraction error: {e}")
            return {}

    def extract_datetime(self, image_path: Path) -> Optional[datetime]:
        """
        Extract the datetime when the photo was taken.

        Returns:
            datetime object or None
        """
        exif_data = self.extract_exif_data(image_path)

        if not exif_data:
            return None

        # Try different datetime tags
        datetime_tags = ['DateTimeOriginal', 'DateTimeDigitized', 'DateTime']

        for tag in datetime_tags:
            if tag in exif_data:
                try:
                    # Parse datetime string (format: "2023:11:26 14:30:00")
                    dt_str = str(exif_data[tag])
                    dt = datetime.strptime(dt_str, "%Y:%m:%d %H:%M:%S")
                    return dt
                except (ValueError, TypeError):
                    continue

        return None

    def extract_gps_coordinates(self, image_path: Path) -> Optional[Tuple[float, float]]:
        """
        Extract GPS coordinates from image EXIF data.

        Returns:
            Tuple of (latitude, longitude) or None
        """
        if not self.metadata_available:
            return None

        try:
            image = Image.open(image_path)
            exif = image._getexif()

            if not exif:
                return None

            # Get GPS info
            gps_info = {}
            for tag_id, value in exif.items():
                tag = TAGS.get(tag_id, tag_id)
                if tag == 'GPSInfo':
                    for gps_tag_id in value:
                        gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                        gps_info[gps_tag] = value[gps_tag_id]

            if not gps_info:
                return None

            # Convert to decimal degrees
            lat = self._convert_to_degrees(gps_info.get('GPSLatitude'))
            lon = self._convert_to_degrees(gps_info.get('GPSLongitude'))

            if lat is None or lon is None:
                return None

            # Adjust for hemisphere
            if gps_info.get('GPSLatitudeRef') == 'S':
                lat = -lat
            if gps_info.get('GPSLongitudeRef') == 'W':
                lon = -lon

            return (lat, lon)

        except Exception as e:
            print(f"  GPS extraction error: {e}")
            return None

    def _convert_to_degrees(self, value) -> Optional[float]:
        """
        Convert GPS coordinates to degrees.

        Args:
            value: GPS coordinate in format ((deg, 1), (min, 1), (sec, 1))

        Returns:
            Decimal degrees or None
        """
        if not value:
            return None

        try:
            d = float(value[0][0]) / float(value[0][1])
            m = float(value[1][0]) / float(value[1][1])
            s = float(value[2][0]) / float(value[2][1])

            return d + (m / 60.0) + (s / 3600.0)
        except (IndexError, TypeError, ZeroDivisionError):
            return None

    def get_location_name(self, coordinates: Tuple[float, float]) -> Optional[str]:
        """
        Get location name from GPS coordinates using reverse geocoding.

        Args:
            coordinates: Tuple of (latitude, longitude)

        Returns:
            Location name (city, state, country) or None
        """
        if not self.geocoder:
            return None

        try:
            lat, lon = coordinates
            location = self.geocoder.reverse(f"{lat}, {lon}", exactly_one=True)

            if location and location.raw.get('address'):
                address = location.raw['address']

                # Try to get city, state, country
                parts = []

                # City
                city = address.get('city') or address.get('town') or address.get('village')
                if city:
                    parts.append(city)

                # State/Region
                state = address.get('state') or address.get('region')
                if state:
                    parts.append(state)

                # Country
                country = address.get('country')
                if country:
                    parts.append(country)

                if parts:
                    return ', '.join(parts)

        except (GeocoderTimedOut, GeocoderServiceError) as e:
            print(f"  Geocoding error: {e}")
        except Exception as e:
            print(f"  Location lookup error: {e}")

        return None

    def get_metadata_summary(self, image_path: Path) -> Dict[str, Any]:
        """
        Get a summary of image metadata.

        Returns:
            Dictionary with datetime, GPS coordinates, and location
        """
        summary = {
            'datetime': None,
            'gps_coordinates': None,
            'location_name': None,
            'year': None,
            'month': None,
            'date_str': None
        }

        # Extract datetime
        dt = self.extract_datetime(image_path)
        if dt:
            summary['datetime'] = dt
            summary['year'] = dt.year
            summary['month'] = dt.month
            summary['date_str'] = dt.strftime("%Y-%m")

        # Extract GPS
        coords = self.extract_gps_coordinates(image_path)
        if coords:
            summary['gps_coordinates'] = coords
            # Get location name (with rate limiting consideration)
            location = self.get_location_name(coords)
            if location:
                summary['location_name'] = location

        return summary


class ImageContentAnalyzer:
    """Analyzes image content using computer vision."""

    def __init__(self):
        """Initialize the image content analyzer."""
        self.vision_available = VISION_AVAILABLE
        self.model = None
        self.processor = None
        self.face_cascade = None

        if self.vision_available:
            try:
                print("Loading CLIP model for image analysis...")
                self.model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
                self.processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

                # Load OpenCV face detection
                cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
                self.face_cascade = cv2.CascadeClassifier(cascade_path)

                print("✓ CLIP model loaded successfully")
            except Exception as e:
                print(f"Warning: Could not load CLIP model: {e}")
                self.vision_available = False

    def detect_people(self, image_path: Path) -> bool:
        """
        Detect if there are people in the image using face detection.

        Returns:
            True if people detected, False otherwise
        """
        if not self.vision_available or self.face_cascade is None:
            return False

        try:
            # Read image
            img = cv2.imread(str(image_path))
            if img is None:
                return False

            # Convert to grayscale for face detection
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

            # Detect faces
            faces = self.face_cascade.detectMultiScale(
                gray,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=(30, 30)
            )

            return len(faces) > 0

        except Exception as e:
            print(f"  Face detection error: {e}")
            return False

    def classify_image_content(self, image_path: Path) -> Dict[str, float]:
        """
        Classify image content using CLIP zero-shot classification.

        Returns:
            Dictionary of category -> confidence score
        """
        if not self.vision_available or self.model is None:
            return {}

        try:
            # Open image
            image = Image.open(image_path)

            # Define categories to check
            categories = [
                "a photo of a home interior room",
                "a photo of a living room",
                "a photo of a bedroom",
                "a photo of a kitchen",
                "a photo of a bathroom",
                "a photo of furniture",
                "a photo of a house exterior",
                "a photo of people",
                "a screenshot",
                "a photo of outdoors",
                "a photo of nature"
            ]

            # Prepare inputs
            inputs = self.processor(
                text=categories,
                images=image,
                return_tensors="pt",
                padding=True
            )

            # Get predictions
            with torch.no_grad():
                outputs = self.model(**inputs)
                logits_per_image = outputs.logits_per_image
                probs = logits_per_image.softmax(dim=1)

            # Convert to dictionary
            scores = {}
            for i, category in enumerate(categories):
                scores[category] = float(probs[0][i])

            return scores

        except Exception as e:
            print(f"  Image classification error: {e}")
            return {}

    def is_home_interior_no_people(self, image_path: Path) -> Tuple[bool, Dict[str, float]]:
        """
        Check if image is a home interior without people.

        Returns:
            Tuple of (is_interior_no_people, classification_scores)
        """
        if not self.vision_available:
            return (False, {})

        # Classify image content
        scores = self.classify_image_content(image_path)

        if not scores:
            return (False, {})

        # Check for home interior indicators
        interior_score = max(
            scores.get("a photo of a home interior room", 0),
            scores.get("a photo of a living room", 0),
            scores.get("a photo of a bedroom", 0),
            scores.get("a photo of a kitchen", 0),
            scores.get("a photo of a bathroom", 0),
            scores.get("a photo of furniture", 0)
        )

        # Check for people
        people_score = scores.get("a photo of people", 0)
        has_faces = self.detect_people(image_path)

        # Determine if it's an interior without people
        is_interior = interior_score > 0.3  # 30% confidence threshold
        has_people = people_score > 0.2 or has_faces  # 20% confidence or face detection

        return (is_interior and not has_people, scores)

    def has_people_in_photo(self, image_path: Path) -> Tuple[bool, Dict[str, float]]:
        """
        Check if image contains people (for social photos).

        Returns:
            Tuple of (has_people, classification_scores)
        """
        if not self.vision_available:
            return (False, {})

        # Classify image content
        scores = self.classify_image_content(image_path)

        if not scores:
            return (False, {})

        # Check for people indicators
        people_score = scores.get("a photo of people", 0)
        has_faces = self.detect_people(image_path)

        # Check that it's NOT a screenshot
        screenshot_score = scores.get("a screenshot", 0)
        is_screenshot = screenshot_score > 0.4  # High threshold for screenshots

        # Determine if photo has people (and is not a screenshot)
        has_people = (people_score > 0.15 or has_faces) and not is_screenshot

        return (has_people, scores)


class ContentBasedFileOrganizer:
    """Organize files based on content analysis using OCR."""

    def __init__(self, base_path: str = None, organize_by_date: bool = False, organize_by_location: bool = False):
        """
        Initialize the organizer.

        Args:
            base_path: Base path for organized files
            organize_by_date: If True, organize photos by date (Photos/2023/11/)
            organize_by_location: If True, organize photos by location when GPS data available
        """
        self.base_path = Path(base_path or os.path.expanduser("~/Documents"))
        self.enricher = MetadataEnricher()
        self.validator = SchemaValidator()
        self.registry = SchemaRegistry()
        self.classifier = ContentClassifier()
        self.image_analyzer = ImageContentAnalyzer()
        self.metadata_parser = ImageMetadataParser()
        self.stats = defaultdict(int)
        self.ocr_available = OCR_AVAILABLE
        self.organize_by_date = organize_by_date
        self.organize_by_location = organize_by_location

        # Filepath-based classification (checked FIRST before content analysis)
        self.filepath_patterns = {
            # Log files
            '.log': 'Technical/Logs',
            '.log.gz': 'Technical/Logs',
            '.out': 'Technical/Logs',

            # Python
            '.py': 'Technical/Python',
            '.pyc': 'Technical/Python/Compiled',
            '.pyw': 'Technical/Python',
            '.pyx': 'Technical/Python',
            '.pyd': 'Technical/Python',

            # JavaScript/TypeScript
            '.js': 'Technical/JavaScript',
            '.jsx': 'Technical/JavaScript',
            '.mjs': 'Technical/JavaScript',
            '.cjs': 'Technical/JavaScript',
            '.ts': 'Technical/TypeScript',
            '.tsx': 'Technical/TypeScript',

            # Web
            '.html': 'Technical/Web',
            '.htm': 'Technical/Web',
            '.css': 'Technical/Web',
            '.scss': 'Technical/Web',
            '.sass': 'Technical/Web',
            '.less': 'Technical/Web',

            # Shell scripts
            '.sh': 'Technical/Shell',
            '.bash': 'Technical/Shell',
            '.zsh': 'Technical/Shell',
            '.fish': 'Technical/Shell',

            # Config files
            '.json': 'Technical/Config',
            '.yaml': 'Technical/Config',
            '.yml': 'Technical/Config',
            '.toml': 'Technical/Config',
            '.ini': 'Technical/Config',
            '.conf': 'Technical/Config',
            '.config': 'Technical/Config',
            '.env': 'Technical/Config',

            # Database
            '.sql': 'Technical/Database',
            '.db': 'Technical/Database',
            '.sqlite': 'Technical/Database',
            '.sqlite3': 'Technical/Database',

            # Java/Kotlin
            '.java': 'Technical/Java',
            '.class': 'Technical/Java/Compiled',
            '.jar': 'Technical/Java/Archives',
            '.kt': 'Technical/Kotlin',
            '.kts': 'Technical/Kotlin',

            # C/C++
            '.c': 'Technical/C',
            '.cpp': 'Technical/C++',
            '.cc': 'Technical/C++',
            '.cxx': 'Technical/C++',
            '.h': 'Technical/C/Headers',
            '.hpp': 'Technical/C++/Headers',

            # Go
            '.go': 'Technical/Go',

            # Rust
            '.rs': 'Technical/Rust',

            # Ruby
            '.rb': 'Technical/Ruby',
            '.rake': 'Technical/Ruby',

            # PHP
            '.php': 'Technical/PHP',

            # Swift
            '.swift': 'Technical/Swift',

            # Markdown and docs
            '.md': 'Technical/Documentation',
            '.markdown': 'Technical/Documentation',
            '.rst': 'Technical/Documentation',
            '.adoc': 'Technical/Documentation',

            # Version control
            '.gitignore': 'Technical/VersionControl',
            '.gitattributes': 'Technical/VersionControl',

            # Build/Package files
            'Makefile': 'Technical/Build',
            'Dockerfile': 'Technical/Build',
            'docker-compose.yml': 'Technical/Build',
            'package.json': 'Technical/Build',
            'package-lock.json': 'Technical/Build',
            'yarn.lock': 'Technical/Build',
            'Cargo.toml': 'Technical/Build',
            'go.mod': 'Technical/Build',
            'requirements.txt': 'Technical/Build',
            'Pipfile': 'Technical/Build',
            'pyproject.toml': 'Technical/Build',
        }

        # Content-based organization structure
        self.category_paths = {
            'legal': {
                'contracts': 'Legal/Contracts',
                'real_estate': 'Legal/RealEstate',
                'corporate': 'Legal/Corporate',
                'other': 'Legal/Other'
            },
            'financial': {
                'tax': 'Financial/Tax',
                'invoices': 'Financial/Invoices',
                'statements': 'Financial/Statements',
                'other': 'Financial/Other'
            },
            'business': {
                'planning': 'Business/Planning',
                'marketing': 'Business/Marketing',
                'proposals': 'Business/Proposals',
                'clients': 'Business/Clients',  # Will be further organized by company name
                'other': 'Business/Other'
            },
            'personal': {
                'employment': 'Personal/Employment',
                'identification': 'Personal/Identification',
                'certificates': 'Personal/Certificates',
                'other': 'Personal/Other'
            },
            'medical': {
                'records': 'Medical/Records',
                'insurance': 'Medical/Insurance',
                'prescriptions': 'Medical/Prescriptions',
                'other': 'Medical/Other'
            },
            'property': {
                'leases': 'Property/Leases',
                'maintenance': 'Property/Maintenance',
                'other': 'Property/Other'
            },
            'education': {
                'coursework': 'Education/Coursework',
                'research': 'Education/Research',
                'records': 'Education/Records',
                'other': 'Education/Other'
            },
            'technical': {
                'documentation': 'Technical/Documentation',
                'architecture': 'Technical/Architecture',
                'other': 'Technical/Other'
            },
            'creative': {
                'design': 'Creative/Design',
                'branding': 'Creative/Branding',
                'photos': 'Creative/Photos',
                'other': 'Creative/Other'
            },
            'property_management': 'Property_Management',
            'game_assets': {
                'audio': 'GameAssets/Audio',
                'music': 'GameAssets/Music',
                'sprites': 'GameAssets/Sprites',
                'textures': 'GameAssets/Textures',
                'other': 'GameAssets/Other'
            },
            'media': {
                'photos': {
                    'screenshots': 'Media/Photos/Screenshots',
                    'travel': 'Media/Photos/Travel',
                    'portraits': 'Media/Photos/Portraits',
                    'events': 'Media/Photos/Events',
                    'documents': 'Media/Photos/Documents',
                    'social': 'Media/Photos/Social',
                    'other': 'Media/Photos/Other'
                },
                'videos': {
                    'recordings': 'Media/Videos/Recordings',
                    'exports': 'Media/Videos/Exports',
                    'screencasts': 'Media/Videos/Screencasts',
                    'other': 'Media/Videos/Other'
                },
                'audio': {
                    'recordings': 'Media/Audio/Recordings',
                    'music': 'Media/Audio/Music',
                    'podcasts': 'Media/Audio/Podcasts',
                    'other': 'Media/Audio/Other'
                }
            },
            'uncategorized': 'Uncategorized'
        }

        # Game asset detection patterns
        self.game_audio_keywords = [
            'bolt', 'spell', 'magic', 'cast', 'chirp', 'crossbow', 'dagger',
            'sword', 'arrow', 'bow', 'heal', 'potion', 'lightning', 'fire',
            'ice', 'acid', 'poison', 'explosion', 'blast', 'summon', 'dispel',
            'petrification', 'neutralize', 'slow', 'darkness', 'achievement',
            'quest', 'unlock', 'lock', 'door', 'chest', 'coin', 'pickup',
            'attack', 'hit', 'damage', 'death', 'footstep', 'jump', 'land',
            'monster', 'creature', 'enemy', 'boss', 'battle', 'combat',
            'starving', 'hunger', 'thirst', 'eat', 'drink', 'sleep',
            'fiddle', 'lute', 'mandoline', 'glockenspiel', 'instrument',
            'identify', 'greater', 'mental'
        ]

        self.game_music_keywords = [
            'battle', 'boss', 'dungeon', 'castle', 'forest', 'town', 'village',
            'temple', 'ruins', 'cave', 'mountain', 'ocean', 'desert', 'snow',
            'victory', 'defeat', 'theme', 'menu', 'credits', 'intro', 'outro',
            'mysterious', 'dark', 'light', 'epic', 'calm', 'peaceful', 'tension',
            'chaos', 'hope', 'despair', 'triumph', 'march', 'symphony', 'monotony',
            'drakalor', 'altar', 'lawful', 'chaotic', 'neutral', 'alignment',
            'dwarven', 'elven', 'orcish', 'halls', 'abandon', 'corrupting',
            'breeze', 'clockwork', 'knowledge', 'oddisey', 'final', 'welcome'
        ]

        self.game_sprite_keywords = [
            'frame', 'item', 'segment', 'sprite', 'texture', 'tile',
            'leg', 'arm', 'head', 'torso', 'body', 'wing', 'tail',
            'hair', 'face', 'eye', 'mouth', 'hand', 'foot',
            'wall', 'floor', 'ceiling', 'door', 'window', 'stairs',
            'tree', 'rock', 'grass', 'water', 'lava', 'cloud',
            'sword', 'shield', 'armor', 'helmet', 'boot', 'glove',
            'potion', 'scroll', 'wand', 'staff', 'ring', 'amulet',
            'coin', 'gem', 'crystal', 'ore', 'metal', 'wood',
            'monster', 'enemy', 'npc', 'character', 'player', 'hero',
            'icon', 'button', 'ui', 'hud', 'menu', 'cursor',
            'particle', 'effect', 'explosion', 'smoke', 'blood',
            'corner', 'edge', 'border', 'container', 'btn', 'talent',
            'bar', 'over', 'down', 'up', 'left', 'right', 'main',
            'bottom', 'top', 'extension', 'descend', 'ascend',
            'mad_carpenter', 'no_more', 'bedroom', 'front', 'back',
            'upper', 'lower', 'middle', 'dead', 'alive', 'sleeping',
            'female', 'male', 'white', 'black', 'red', 'blue', 'green',
            'silver', 'gold', 'bronze', 'iron', 'steel', 'mithril',
            'hills', 'road', 'path', 'bridge', 'gate', 'fence',
            # Additional game environment and UI keywords
            'tentacle', 'shadow', 'altar', 'dungeon', 'throne', 'torch',
            'cloak', 'champion', 'curse', 'warning', 'mouse', 'slider',
            'decal', 'column', 'banner', 'sewer', 'statue', 'pillar',
            'orc', 'dwarf', 'elf', 'hurth', 'helf', 'troll', 'goblin',
            'fire', 'ice', 'sand', 'mount', 'tmount', 'deco', 'entrance',
            'pupils', 'shoulders', 'stunned', 'poisoned', 'blind', 'deaf',
            'slowed', 'levitating', 'hungry', 'strained', 'next', 'prev',
            'groove', 'handle', 'cube', 'psf', 'inventory'
        ]

    def classify_by_filepath(self, file_path: Path) -> Optional[str]:
        """
        Classify file based on filepath patterns (extension, filename).

        Returns:
            Category path string if matched, None otherwise
        """
        # Check exact filename matches first (e.g., Makefile, Dockerfile)
        filename = file_path.name
        if filename in self.filepath_patterns:
            return self.filepath_patterns[filename]

        # Check file extension
        ext = file_path.suffix.lower()
        if ext in self.filepath_patterns:
            base_path = self.filepath_patterns[ext]

            # Try to extract project name from path
            project_name = self.extract_project_name(file_path)
            if project_name:
                # Add project subdirectory (e.g., Technical/Python/MyProject)
                return f"{base_path}/{project_name}"

            return base_path

        # Check double extensions (e.g., .log.gz)
        if len(file_path.suffixes) >= 2:
            double_ext = ''.join(file_path.suffixes[-2:]).lower()
            if double_ext in self.filepath_patterns:
                return self.filepath_patterns[double_ext]

        return None

    def extract_project_name(self, file_path: Path) -> Optional[str]:
        """
        Extract project name from file path.

        Looks for common project indicators in path:
        - Directory names like 'myproject', 'my-app', etc.
        - Skips common non-project directories

        Returns:
            Project name if found, None otherwise
        """
        skip_dirs = {
            'src', 'lib', 'bin', 'dist', 'build', 'out', 'target',
            'node_modules', 'venv', '.venv', 'env', '__pycache__',
            'scripts', 'tests', 'test', 'docs', 'doc', 'examples',
            'static', 'public', 'assets', 'resources', 'config',
            'home', 'users', 'documents', 'downloads', 'desktop',
            'code', 'projects', 'dev', 'work', 'repos', 'git'
        }

        # Get all parent directories
        parts = file_path.parts

        # Look backwards from the file for a likely project directory
        for i in range(len(parts) - 2, -1, -1):  # Skip the filename itself
            dir_name = parts[i].lower()

            # Skip common non-project directories
            if dir_name in skip_dirs:
                continue

            # Skip hidden directories
            if dir_name.startswith('.'):
                continue

            # Found a likely project directory
            # Return with original case preserved
            return parts[i]

        return None

    def classify_game_asset(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Classify file as a game asset based on filename patterns.

        Returns:
            Tuple of (category, subcategory) or None if not a game asset
        """
        filename = file_path.name.lower()
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()

        # Check for audio files (.wav, .ogg, .mp3)
        if ext in ['.wav', '.ogg', '.mp3', '.flac', '.aac']:
            # Check for game music patterns (usually .ogg files with specific names)
            if ext == '.ogg':
                for keyword in self.game_music_keywords:
                    if keyword in stem:
                        return ('game_assets', 'music')

            # Check for game sound effects
            for keyword in self.game_audio_keywords:
                if keyword in stem:
                    return ('game_assets', 'audio')

        # Check for image files that are game sprites/textures
        if ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tga', '.dds']:
            # Check for sprite/texture patterns
            for keyword in self.game_sprite_keywords:
                if keyword in stem:
                    # Distinguish between sprites and textures
                    if any(kw in stem for kw in ['frame', 'sprite', 'leg', 'arm', 'head', 'torso', 'body', 'wing', 'hair', 'face', 'mouth']):
                        return ('game_assets', 'sprites')
                    else:
                        return ('game_assets', 'textures')

        return None

    def classify_media_file(self, file_path: Path, image_metadata: Dict = None) -> Optional[Tuple[str, str, str]]:
        """
        Classify media files (photos, videos, audio) into subcategories.

        Returns:
            Tuple of (category, media_type, subcategory) or None if not a media file
            Example: ('media', 'photos', 'screenshots') or ('media', 'videos', 'recordings')
        """
        filename = file_path.name.lower()
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()

        # Videos - .mp4, .mov, .avi, .mkv, .webm, .m4v
        if ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.m4v', '.flv', '.wmv']:
            # Screen recordings
            if 'screen' in stem or 'recording' in stem or 'capture' in stem:
                return ('media', 'videos', 'screencasts')
            # Exports (from video editors)
            elif 'export' in stem or 'render' in stem or 'final' in stem or 'cut' in stem:
                return ('media', 'videos', 'exports')
            # Default to recordings
            else:
                return ('media', 'videos', 'recordings')

        # Audio - .mp3, .wav, .m4a, .aac, .flac, .ogg (but not game music)
        if ext in ['.mp3', '.m4a', '.aac', '.flac', '.wma']:
            # Podcasts
            if 'podcast' in stem or 'episode' in stem or 'interview' in stem:
                return ('media', 'audio', 'podcasts')
            # Music
            elif 'song' in stem or 'album' in stem or 'track' in stem or 'music' in stem:
                return ('media', 'audio', 'music')
            # Voice recordings
            elif 'recording' in stem or 'voice' in stem or 'memo' in stem or 'audio' in stem:
                return ('media', 'audio', 'recordings')
            # Default to recordings
            else:
                return ('media', 'audio', 'recordings')

        # Photos - .jpg, .jpeg, .png, .heic, .gif, .webp, .bmp
        if ext in ['.jpg', '.jpeg', '.png', '.heic', '.gif', '.webp', '.bmp', '.tiff', '.tif']:
            # Screenshots (highest priority for photos)
            if filename.startswith('screenshot') or 'screen shot' in filename:
                return ('media', 'photos', 'screenshots')

            # Scanned documents/receipts (OCR will detect text)
            if 'scan' in stem or 'receipt' in stem or 'document' in stem or 'invoice' in stem:
                return ('media', 'photos', 'documents')

            # Travel photos (has GPS metadata)
            if image_metadata and image_metadata.get('gps_coordinates'):
                # If we have GPS coordinates, it's likely a travel photo
                return ('media', 'photos', 'travel')

            # Photos with datetime (camera photos) - organize by type
            if image_metadata and image_metadata.get('datetime'):
                # Photos with camera EXIF data are likely personal photos
                # Default to 'other' category for general photos
                return ('media', 'photos', 'other')

            # Photos without metadata - still categorize as media if they're actual photos
            # (as opposed to game sprites which would be caught earlier)
            if ext in ['.jpg', '.jpeg', '.heic']:
                return ('media', 'photos', 'other')

            # PNG files without clear classification fall through
            # (could be screenshots, documents, or game assets that weren't caught)
            return None

        return None

    def extract_text_from_image(self, image_path: Path) -> str:
        """Extract text from image using OCR."""
        if not self.ocr_available:
            return ""

        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            print(f"  OCR error: {e}")
            return ""

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF (searchable or scanned)."""
        if not self.ocr_available:
            return ""

        text = ""

        try:
            # First try to extract text directly (for searchable PDFs)
            with open(pdf_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages[:10]:  # Limit to first 10 pages
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            # If we got meaningful text, return it
            if len(text.strip()) > 100:
                return text.strip()

            # Otherwise, try OCR on the PDF
            print(f"  Using OCR for scanned PDF...")
            images = convert_from_path(pdf_path, first_page=1, last_page=5)
            for image in images:
                text += pytesseract.image_to_string(image) + "\n"

            return text.strip()
        except Exception as e:
            print(f"  PDF extraction error: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            return ""

        try:
            doc = Document(docx_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)
        except Exception as e:
            print(f"  DOCX extraction error: {e}")
            return ""

    def extract_text_from_xlsx(self, xlsx_path: Path) -> str:
        """Extract text from Excel spreadsheet."""
        if not EXCEL_AVAILABLE:
            return ""

        try:
            workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
            text = []

            # Limit to first 5 sheets
            for sheet_name in list(workbook.sheetnames)[:5]:
                sheet = workbook[sheet_name]
                # Limit to first 100 rows
                for row in list(sheet.iter_rows(max_row=100, values_only=True)):
                    row_text = ' '.join([str(cell) for cell in row if cell is not None])
                    if row_text.strip():
                        text.append(row_text)

            workbook.close()
            return "\n".join(text)
        except Exception as e:
            print(f"  XLSX extraction error: {e}")
            return ""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from various file types."""
        mime_type = self.enricher.detect_mime_type(str(file_path))
        file_ext = file_path.suffix.lower()

        # Images
        if mime_type and mime_type.startswith('image/'):
            return self.extract_text_from_image(file_path)

        # PDFs
        elif mime_type == 'application/pdf' or file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)

        # Word documents
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)

        # Excel spreadsheets
        elif file_ext in ['.xlsx', '.xls']:
            return self.extract_text_from_xlsx(file_path)

        # Text files
        elif mime_type and mime_type.startswith('text/') or file_ext in ['.txt', '.md', '.csv']:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(50000)  # First 50KB
            except Exception:
                return ""

        return ""

    def detect_file_category(self, file_path: Path) -> Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]:
        """
        Detect file category based on content.

        Priority order:
        1. Game asset detection (audio, sprites, textures)
        2. Filepath-based classification (file extensions, filenames)
        3. Image content analysis (for home interiors)
        4. OCR and text-based classification

        Returns:
            Tuple of (main_category, subcategory, schema_type, extracted_text, company_name, people_names, image_metadata)
        """
        # PRIORITY 0: Check for game assets FIRST (before filepath patterns)
        game_asset = self.classify_game_asset(file_path)
        if game_asset:
            category, subcategory = game_asset
            print(f"  ✓ Game asset detected: {subcategory}")

            # Determine schema type
            mime_type = self.enricher.detect_mime_type(str(file_path))
            if mime_type:
                if mime_type.startswith('image/'):
                    schema_type = 'ImageObject'
                elif mime_type.startswith('audio/'):
                    schema_type = 'AudioObject'
                elif mime_type.startswith('video/'):
                    schema_type = 'VideoObject'
                else:
                    schema_type = 'DigitalDocument'
            else:
                schema_type = 'DigitalDocument'

            return (category, subcategory, schema_type, '', None, [], {})

        # PRIORITY 1: Check filepath patterns (most efficient and accurate for code files)
        filepath_category = self.classify_by_filepath(file_path)
        if filepath_category:
            print(f"  ✓ Filepath match: {filepath_category}")
            # Still need to determine schema type
            mime_type = self.enricher.detect_mime_type(str(file_path))
            if mime_type:
                if mime_type.startswith('image/'):
                    schema_type = 'ImageObject'
                elif mime_type == 'application/pdf':
                    schema_type = 'DigitalDocument'
                elif mime_type.startswith('video/'):
                    schema_type = 'VideoObject'
                elif mime_type.startswith('audio/'):
                    schema_type = 'AudioObject'
                else:
                    schema_type = 'DigitalDocument'
            else:
                schema_type = 'DigitalDocument'

            # Return filepath-based category as a special marker
            # We'll handle this in get_destination_path
            return ('filepath', filepath_category, schema_type, '', None, [], {})

        # Determine schema type for non-filepath matches
        mime_type = self.enricher.detect_mime_type(str(file_path))
        if mime_type:
            if mime_type.startswith('image/'):
                schema_type = 'ImageObject'
            elif mime_type == 'application/pdf':
                schema_type = 'DigitalDocument'
            elif mime_type.startswith('video/'):
                schema_type = 'VideoObject'
            elif mime_type.startswith('audio/'):
                schema_type = 'AudioObject'
            else:
                schema_type = 'DigitalDocument'
        else:
            schema_type = 'DigitalDocument'

        # Extract metadata for images
        image_metadata = {}
        if schema_type == 'ImageObject' and self.metadata_parser.metadata_available:
            print(f"  Extracting image metadata...")
            image_metadata = self.metadata_parser.get_metadata_summary(file_path)

            if image_metadata.get('datetime'):
                dt = image_metadata['datetime']
                print(f"  ✓ Photo taken: {dt.strftime('%Y-%m-%d %H:%M:%S')}")

            if image_metadata.get('gps_coordinates'):
                coords = image_metadata['gps_coordinates']
                print(f"  ✓ GPS: {coords[0]:.6f}, {coords[1]:.6f}")

            if image_metadata.get('location_name'):
                print(f"  ✓ Location: {image_metadata['location_name']}")

        # PRIORITY 1.5: Check for media files (photos, videos, audio)
        # This runs after metadata extraction so we can use GPS/datetime for classification
        media_classification = self.classify_media_file(file_path, image_metadata)
        if media_classification:
            category, media_type, subcategory = media_classification
            print(f"  ✓ Media file detected: {media_type}/{subcategory}")
            return (category, f"{media_type}_{subcategory}", schema_type, '', None, [], image_metadata)

        # PRIORITY 2: Check for photos with people (social)
        if schema_type == 'ImageObject' and self.image_analyzer.vision_available:
            print(f"  Analyzing image content...")

            # First check if photo has people
            has_people, people_scores = self.image_analyzer.has_people_in_photo(file_path)

            if has_people:
                print(f"  ✓ Detected: Photo with people")
                if people_scores:
                    top_categories = sorted(people_scores.items(), key=lambda x: x[1], reverse=True)[:3]
                    print(f"  Top matches: {', '.join([f'{cat}: {score:.2%}' for cat, score in top_categories])}")
                return ('media', 'photos_social', schema_type, '', None, [], image_metadata)

            # Then check for home interior without people
            is_property_mgmt, vision_scores = self.image_analyzer.is_home_interior_no_people(file_path)

            if is_property_mgmt:
                print(f"  ✓ Detected: Home interior without people")
                if vision_scores:
                    top_categories = sorted(vision_scores.items(), key=lambda x: x[1], reverse=True)[:3]
                    print(f"  Top matches: {', '.join([f'{cat}: {score:.2%}' for cat, score in top_categories])}")
                return ('property_management', 'other', schema_type, '', None, [], image_metadata)

        # PRIORITY 3: Regular text extraction and classification
        print(f"  Extracting content...")
        extracted_text = self.extract_text(file_path)

        if extracted_text:
            print(f"  Extracted {len(extracted_text)} characters")
            category, subcategory, company_name, people_names = self.classifier.classify_content(extracted_text, file_path.name)
            if company_name:
                print(f"  Detected company: {company_name}")
            if people_names:
                print(f"  Detected people: {', '.join(people_names[:3])}{' ...' if len(people_names) > 3 else ''}")
            print(f"  Classified as: {category}/{subcategory}")
        else:
            print(f"  No text extracted, using filename")
            category, subcategory, company_name, people_names = self.classifier.classify_content("", file_path.name)

        return (category, subcategory, schema_type, extracted_text, company_name, people_names, image_metadata)

    def generate_schema(self, file_path: Path, schema_type: str, extracted_text: str = "") -> Dict:
        """Generate Schema.org metadata for a file with extracted content."""
        stats = file_path.stat()
        mime_type = self.enricher.detect_mime_type(str(file_path))
        file_url = f"https://localhost/files/{quote(file_path.name)}"
        actual_path = str(file_path.absolute())

        # Create generator based on type
        if schema_type == 'ImageObject':
            generator = ImageGenerator(schema_type)
            generator.set_basic_info(
                name=file_path.name,
                content_url=file_url,
                encoding_format=mime_type or 'image/png',
                description=f"{file_path.name}"
            )
        elif schema_type in ['DigitalDocument', 'Article']:
            generator = DocumentGenerator(schema_type)
            generator.set_basic_info(
                name=file_path.name,
                description=f"{file_path.name}"
            )
            generator.set_file_info(
                encoding_format=mime_type or 'application/octet-stream',
                url=file_url,
                content_size=stats.st_size
            )
        else:
            generator = DocumentGenerator()
            generator.set_basic_info(
                name=file_path.name,
                description=f"{file_path.name}"
            )

        # Set dates
        try:
            generator.set_dates(
                created=datetime.fromtimestamp(stats.st_ctime),
                modified=datetime.fromtimestamp(stats.st_mtime)
            )
        except Exception:
            pass

        # Add extracted text as abstract/text property
        if extracted_text:
            try:
                # Truncate to reasonable length for schema
                text_preview = extracted_text[:1000] + ('...' if len(extracted_text) > 1000 else '')
                generator.set_property('abstract', text_preview, PropertyType.TEXT)
                generator.set_property('text', extracted_text[:5000], PropertyType.TEXT)
            except Exception:
                pass

        # Add file path
        try:
            generator.set_property('filePath', actual_path, PropertyType.TEXT)
        except Exception:
            pass

        return generator.to_dict()

    def get_destination_path(self, file_path: Path, category: str, subcategory: str, company_name: Optional[str] = None, image_metadata: Optional[Dict] = None) -> Path:
        """
        Get the destination path for a file based on content category.

        Args:
            file_path: Path to the file
            category: Main category
            subcategory: Subcategory
            company_name: Optional company name for business files
            image_metadata: Optional metadata for images (datetime, location)

        Returns:
            Destination path for the file
        """
        # Special handling for filepath-based classification
        if category == 'filepath':
            # subcategory contains the full path (e.g., "Technical/Python/MyProject")
            relative_path = subcategory
        # Special handling for media files with nested structure
        elif category == 'media' and '_' in subcategory:
            # subcategory format: "media_type_subcategory" (e.g., "photos_screenshots")
            parts = subcategory.split('_', 1)  # Split into at most 2 parts
            if len(parts) == 2:
                media_type, media_subcat = parts
                if media_type in self.category_paths['media']:
                    media_dict = self.category_paths['media'][media_type]
                    if isinstance(media_dict, dict):
                        relative_path = media_dict.get(media_subcat, media_dict.get('other', f'Media/{media_type.capitalize()}/Other'))
                    else:
                        relative_path = media_dict
                else:
                    relative_path = 'Media/Other'
            else:
                relative_path = 'Media/Other'
        elif category in self.category_paths:
            if isinstance(self.category_paths[category], dict):
                if subcategory in self.category_paths[category]:
                    relative_path = self.category_paths[category][subcategory]
                else:
                    relative_path = self.category_paths[category].get('other', f'{category.capitalize()}/Other')
            else:
                relative_path = self.category_paths[category]
        else:
            relative_path = 'Uncategorized'

        # If it's a client file with a company name, create company-specific folder
        if subcategory == 'clients' and company_name:
            sanitized_company = self.classifier.sanitize_company_name(company_name)
            relative_path = f"{relative_path}/{sanitized_company}"

        # Date-based organization for images (if enabled and metadata available)
        if self.organize_by_date and image_metadata and image_metadata.get('year'):
            year = image_metadata['year']
            month = image_metadata['month']
            relative_path = f"Photos/{year}/{month:02d}"

        # Location-based organization for images (if enabled and location available)
        elif self.organize_by_location and image_metadata and image_metadata.get('location_name'):
            # Clean location name for folder
            location = image_metadata['location_name']
            # Take first part (usually city)
            city = location.split(',')[0].strip()
            # Sanitize for folder name
            safe_city = re.sub(r'[<>:"/\\|?*]', '', city)
            relative_path = f"Photos/Locations/{safe_city}"

        dest_dir = self.base_path / relative_path
        dest_dir.mkdir(parents=True, exist_ok=True)

        # Handle duplicate filenames
        dest_path = dest_dir / file_path.name
        if dest_path.exists() and dest_path != file_path:
            stem = file_path.stem
            suffix = file_path.suffix
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dest_path = dest_dir / f"{stem}_{timestamp}{suffix}"

        return dest_path

    def should_skip_file(self, file_path: Path) -> bool:
        """Check if file should be skipped."""
        skip_files = {'.DS_Store', '.localized', 'Thumbs.db', 'desktop.ini'}
        skip_dirs = {'__pycache__', '.git', 'node_modules', '.venv', 'venv'}

        if file_path.name.startswith('.') and file_path.name not in {'.gitignore', '.env.example'}:
            return True

        if file_path.name in skip_files:
            return True

        if any(skip_dir in file_path.parts for skip_dir in skip_dirs):
            return True

        return False

    def organize_file(self, file_path: Path, dry_run: bool = False) -> Dict:
        """
        Organize a single file based on content.

        Args:
            file_path: Path to the file
            dry_run: If True, don't actually move files

        Returns:
            Dictionary with organization details
        """
        result = {
            'source': str(file_path),
            'status': 'skipped',
            'reason': None,
            'destination': None,
            'schema': None,
            'extracted_text_length': 0
        }

        if self.should_skip_file(file_path):
            result['reason'] = 'system_file'
            self.stats['skipped'] += 1
            return result

        if not file_path.is_file():
            result['reason'] = 'not_file'
            self.stats['skipped'] += 1
            return result

        try:
            # Detect category based on content
            category, subcategory, schema_type, extracted_text, company_name, people_names, image_metadata = self.detect_file_category(file_path)
            result['extracted_text_length'] = len(extracted_text)
            result['company_name'] = company_name
            result['people_names'] = people_names
            result['image_metadata'] = image_metadata

            # Generate schema with extracted content
            schema = self.generate_schema(file_path, schema_type, extracted_text)

            # Validate schema
            validation_report = self.validator.validate(schema)

            # Get destination path (with optional date/location organization for images)
            dest_path = self.get_destination_path(file_path, category, subcategory, company_name, image_metadata)

            # Skip if already in the right place
            if file_path == dest_path:
                result['status'] = 'already_organized'
                result['destination'] = str(dest_path)
                result['schema'] = schema
                result['category'] = category
                result['subcategory'] = subcategory
                self.stats['already_organized'] += 1
                return result

            # Move file if not dry run
            if not dry_run:
                shutil.move(str(file_path), str(dest_path))

                # Register schema
                schema['url'] = f"file://{dest_path.absolute()}"
                metadata = {
                    'category': category,
                    'subcategory': subcategory,
                    'organized_date': datetime.now().isoformat(),
                    'is_valid': validation_report.is_valid(),
                    'has_extracted_text': bool(extracted_text)
                }
                if company_name:
                    metadata['company_name'] = company_name

                self.registry.register(
                    str(dest_path),
                    schema,
                    metadata=metadata
                )

            result['status'] = 'organized' if not dry_run else 'would_organize'
            result['destination'] = str(dest_path)
            result['schema'] = schema
            result['category'] = category
            result['subcategory'] = subcategory
            result['is_valid'] = validation_report.is_valid()

            self.stats['organized'] += 1
            self.stats[f'{category}_{subcategory}'] += 1

        except Exception as e:
            result['status'] = 'error'
            result['reason'] = str(e)
            self.stats['errors'] += 1
            print(f"  ✗ Error: {e}")

        return result

    def scan_directory(self, directory: Path) -> List[Path]:
        """Scan directory for files to organize."""
        files = []
        try:
            for item in directory.rglob('*'):
                if item.is_file() and not self.should_skip_file(item):
                    files.append(item)
        except PermissionError:
            print(f"Permission denied: {directory}")
        return files

    def organize_directories(self, source_dirs: List[str], dry_run: bool = False, limit: int = None) -> Dict:
        """
        Organize files from multiple source directories.

        Args:
            source_dirs: List of directory paths to organize
            dry_run: If True, simulate organization without moving files
            limit: Maximum number of files to process (for testing)

        Returns:
            Dictionary with organization results
        """
        results = []

        print(f"\n{'='*60}")
        print(f"Content-Based File Organization {'(DRY RUN)' if dry_run else ''}")
        print(f"{'='*60}\n")

        if not self.ocr_available:
            print("⚠️  WARNING: OCR libraries not available")
            print("   Install with: pip install pytesseract Pillow pypdf pdf2image")
            print("   Content classification will be limited to filenames\n")

        # Scan all directories
        all_files = []
        for source_dir in source_dirs:
            source_path = Path(source_dir).expanduser()
            if source_path.exists():
                print(f"Scanning: {source_path}")
                files = self.scan_directory(source_path)
                all_files.extend(files)
                print(f"  Found {len(files)} files")
            else:
                print(f"Directory not found: {source_path}")

        if limit:
            all_files = all_files[:limit]
            print(f"\n⚠️  Processing limited to first {limit} files for testing\n")

        print(f"\nTotal files to process: {len(all_files)}\n")

        # Organize each file
        for i, file_path in enumerate(all_files, 1):
            print(f"[{i}/{len(all_files)}] Processing: {file_path.name}")
            result = self.organize_file(file_path, dry_run=dry_run)
            results.append(result)

            if result['status'] == 'organized' or result['status'] == 'would_organize':
                print(f"  → {result['destination']}")
            elif result['status'] == 'error':
                print(f"  ✗ Error: {result['reason']}")

        # Generate summary
        summary = {
            'total_files': len(all_files),
            'organized': self.stats['organized'],
            'already_organized': self.stats['already_organized'],
            'skipped': self.stats['skipped'],
            'errors': self.stats['errors'],
            'dry_run': dry_run,
            'results': results,
            'registry_stats': self.registry.get_statistics() if not dry_run else None
        }

        return summary

    def print_summary(self, summary: Dict):
        """Print organization summary."""
        print(f"\n{'='*60}")
        print("Organization Summary")
        print(f"{'='*60}\n")

        print(f"Total files processed: {summary['total_files']}")
        print(f"Successfully organized: {summary['organized']}")
        print(f"Already organized: {summary['already_organized']}")
        print(f"Skipped: {summary['skipped']}")
        print(f"Errors: {summary['errors']}")

        if summary['dry_run']:
            print("\n⚠️  This was a DRY RUN - no files were moved")

        # Category breakdown
        print(f"\n{'='*60}")
        print("Category Breakdown")
        print(f"{'='*60}\n")

        category_stats = defaultdict(int)
        for result in summary['results']:
            if result.get('category'):
                category_stats[result['category']] += 1

        for category, count in sorted(category_stats.items()):
            print(f"{category.capitalize()}: {count} files")

        # OCR stats
        ocr_count = sum(1 for r in summary['results'] if r.get('extracted_text_length', 0) > 0)
        print(f"\n{'='*60}")
        print("Content Extraction Stats")
        print(f"{'='*60}\n")
        print(f"Files with extracted text: {ocr_count}/{summary['total_files']}")

        # Company detection stats
        company_files = [r for r in summary['results'] if r.get('company_name')]
        if company_files:
            print(f"\n{'='*60}")
            print("Detected Companies")
            print(f"{'='*60}\n")
            company_counts = defaultdict(int)
            for result in company_files:
                company_counts[result['company_name']] += 1

            print(f"Total files with detected companies: {len(company_files)}")
            print(f"\nCompanies found:")
            for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
                print(f"  {company}: {count} files")

        if summary.get('registry_stats'):
            print(f"\n{'='*60}")
            print("Schema Registry")
            print(f"{'='*60}\n")
            stats = summary['registry_stats']
            print(f"Total schemas: {stats['total_schemas']}")
            print(f"Types: {', '.join(stats['types'])}")

    def save_report(self, summary: Dict, output_path: str = None):
        """Save detailed organization report to JSON."""
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"results/content_organization_report_{timestamp}.json"

        with open(output_path, 'w') as f:
            json.dump(summary, f, indent=2, default=str)

        print(f"\nDetailed report saved to: {output_path}")


def main():
    """Main entry point."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Organize files by content using OCR and Schema.org metadata'
    )
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Simulate organization without moving files'
    )
    parser.add_argument(
        '--base-path',
        default='~/Documents',
        help='Base path for organized files (default: ~/Documents)'
    )
    parser.add_argument(
        '--sources',
        nargs='+',
        default=['~/Desktop', '~/Downloads'],
        help='Source directories to organize (default: ~/Desktop ~/Downloads)'
    )
    parser.add_argument(
        '--report',
        help='Path to save detailed JSON report'
    )
    parser.add_argument(
        '--limit',
        type=int,
        help='Limit number of files to process (for testing)'
    )

    args = parser.parse_args()

    # Create organizer
    organizer = ContentBasedFileOrganizer(base_path=args.base_path)

    # Organize directories
    summary = organizer.organize_directories(
        source_dirs=args.sources,
        dry_run=args.dry_run,
        limit=args.limit
    )

    # Print summary
    organizer.print_summary(summary)

    # Save report
    if args.report or not args.dry_run:
        organizer.save_report(summary, args.report)


if __name__ == '__main__':
    main()
